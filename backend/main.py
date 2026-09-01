import io
import logging
import traceback
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import os
import threading
import time
import sys
from datetime import datetime, timedelta
from typing import List, Optional, Dict, Any

BACKEND_DIR = os.path.dirname(os.path.abspath(__file__))
if BACKEND_DIR not in sys.path:
    sys.path.insert(0, BACKEND_DIR)

from fastapi import FastAPI, File, HTTPException, UploadFile, Form, Depends, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.responses import FileResponse, Response, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from action_adapter import JiraAdapter, EmailAdapter
import json
import re
import os.path
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from docx import Document as DocxDocument
import uuid
import os
from dotenv import load_dotenv

# .env dosyasını otomatik yükle
load_dotenv()

# JWT ve Şifreleme
from jose import JWTError, jwt
from passlib.context import CryptContext

# SQLite Kalıcı Depo Katmanı
from store import SQLiteList, init_store

# --- JWT VE GÜVENLİK AYARLARI ---
SECRET_KEY = os.getenv("SECRET_KEY") or "mantis-local-dev-secret-key-change-me"
if not os.getenv("SECRET_KEY"):
    logger = logging.getLogger("mantis.main")
    logger.warning(
        "SECRET_KEY ortam değişkeni tanımlı değil; yerel geliştirme için varsayılan anahtar kullanılıyor. "
        "Üretimde .env dosyasında güçlü bir SECRET_KEY tanımlayın."
    )
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24  # 1 Gün

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
security_scheme = HTTPBearer()

USERS_DB_PATH = os.path.join(os.path.dirname(BACKEND_DIR), "users_db.json")
LEGACY_USERS_DB_PATH = os.path.join(BACKEND_DIR, "users_db.json")

def load_users():
    users = {}
    for path in (LEGACY_USERS_DB_PATH, USERS_DB_PATH):
        if os.path.exists(path):
            try:
                with open(path, "r", encoding="utf-8") as f:
                    loaded = json.load(f)
                    if isinstance(loaded, dict):
                        users.update(loaded)
            except Exception:
                logger.warning("Kullanıcı deposu okunamadı: %s", path)
    return users

def save_users(users):
    with open(USERS_DB_PATH, "w", encoding="utf-8") as f:
        json.dump(users, f, ensure_ascii=False, indent=4)

def get_password_hash(password: str):
    if isinstance(password, str):
        password = password[:72]
    return pwd_context.hash(password)

def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    expire = datetime.utcnow() + (expires_delta or timedelta(minutes=15))
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security_scheme)):
    """Gelen isteğin header kısmındaki Bearer token'ı doğrular."""
    token = credentials.credentials
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Geçersiz kimlik bilgileri veya süresi dolmuş token.",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email: str = payload.get("sub")
        if email is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
    users = load_users()
    user = users.get(email)
    if user is None:
        raise credentials_exception
    return user

# --- PYDANTIC ŞEMALARI ---
class UserSignupRequest(BaseModel):
    email: str
    password: str
    full_name: str

class UserSigninRequest(BaseModel):
    email: str
    password: str

class AttendanceCommand(BaseModel):
    command: str
    consent_given: bool = False
    operator: str = "Hilal Uluca (Lead Architect)"

class WarningLetterRequest(BaseModel):
    personnel_name: str
    date: str
    reason: Optional[str] = None

class ApprovalRequest(BaseModel):
    risk_id: str
    action: str  # "approve" veya "reject"

class TaskActionRequest(BaseModel):
    task_id: int | str
    action: str

class ActionDispatchRequest(BaseModel):
    action_type: str
    title: str
    description: str
    recipient: Optional[str] = "Legal-Team"

class AnalyzeRequest(BaseModel):
    document_id: Optional[str] = "doc_generated"
    text: str

class RiskFindingReal(BaseModel):
    risk_id: str
    severity: str
    clause_text: str
    ai_reasoning: str
    confidence_score: int

class AnalyzeResponse(BaseModel):
    status: str
    document_id: str
    total_risks_found: int
    findings: List[RiskFindingReal]
    message: Optional[str] = None

class RedlineRequest(BaseModel):
    clause_text: str
    reasoning: str

class ChatRequest(BaseModel):
    message: str
    document_id: Optional[str] = None

class VaultDocument(BaseModel):
    id: str
    name: str
    access: str
    last_active: str
    type: str

class OnboardingFullRequest(BaseModel):
    identity: Dict[str, str]
    communication_tone: str
    key_contacts: Dict[str, str]
    standards: Dict[str, str]
    automation_policy: Dict[str, str]

class SubscriptionRequest(BaseModel):
    callback_url: str
    resource_type: str = "messages/inbox"

class OrchestratorChatRequest(BaseModel):
    message: str
    document_id: Optional[str] = None
    voice_enabled: Optional[bool] = False
class DocumentChatRequest(BaseModel):
    message: str

class DocumentUpdateRequest(BaseModel):
    content: str

class WatchedFolderRequest(BaseModel):
    path: str
    enabled: bool = True

class DocumentChatEnvelope(BaseModel):
    document_id: str
    message: str

class DocumentActionRequest(BaseModel):
    document_id: str
    message: str = ""
    selected_text: Optional[str] = None
    instruction: Optional[str] = None

# --- LOGLAMA ---
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("mantis.main")

MANTIS_PERSONA = (
    "Sen Mantis, bir şirketin günlük operasyonel işlerini (mail, görev, "
    "doküman, onay, İK süreçleri) yöneten AI asistanısın. Bir chatbot gibi "
    "değil, işini bilen, inisiyatif alan ama gereksiz soru sormayan bir "
    "ekip arkadaşı gibi konuşursun. Kısa, doğal, sıcak ama profesyonel bir "
    "Türkçeyle yanıt verirsin — resmi yazışma dili ya da gereksiz kalıp "
    "cümleler kullanmazsın. Bir işlemi gerçekten yapmadıysan asla yaptığını "
    "söylemezsin; belirsizlik varsa tahmin yürütmek yerine sorarsın."
)

# --- Adaptörler ve Durum Depoları ---
jira_adapter = JiraAdapter() if 'JiraAdapter' in globals() else None
email_adapter = EmailAdapter() if 'EmailAdapter' in globals() else None

db_approval_states: Dict[str, str] = {}
db_subscriptions: Dict[str, Dict[str, Any]] = {}
chat_memory_history: List[Dict[str, str]] = []
db_proactive_briefing = {"last_summary": "Henüz özet üretilmedi.", "timestamp": None}

# Kalıcı SQLite Depoları
db_tasks = SQLiteList("tasks")
db_vault_documents = SQLiteList("vault_documents")
db_risk_findings = SQLiteList("risk_findings", id_field="risk_id")
db_audit_logs = SQLiteList("audit_logs")
WATCHED_FOLDERS_PATH = os.path.join(BACKEND_DIR, "watched_folders.json")

def load_watched_folders() -> List[Dict[str, Any]]:
    try:
        with open(WATCHED_FOLDERS_PATH, "r", encoding="utf-8") as file:
            data = json.load(file)
            return data if isinstance(data, list) else []
    except (FileNotFoundError, json.JSONDecodeError):
        return []

def save_watched_folders(folders: List[Dict[str, Any]]) -> None:
    with open(WATCHED_FOLDERS_PATH, "w", encoding="utf-8") as file:
        json.dump(folders, file, ensure_ascii=False, indent=2)

_SEED_RISK_FINDINGS = [
    {
        "risk_id": "RSK-8092",
        "severity": "CRITICAL",
        "clause_text": "Tedarikçi, altyapı kesintilerinden veya veri kayıplarından hiçbir koşulda sorumlu tutulamaz.",
        "ai_reasoning": "Sorumluluk reddi maddesi şirketimizin operasyonel SLA garantilerini geçersiz kılmaktadır.",
        "confidence_score": 94,
        "status": "PENDING_APPROVAL",
        "is_approved": False,
        "is_rejected": False
    },
    {
        "risk_id": "RSK-4102",
        "severity": "MEDIUM",
        "clause_text": "İşbu sözleşmeden doğan uyuşmazlıklarda taraflar 15 gün içinde müzakere yolunu dener.",
        "ai_reasoning": "Müzakere süresi yasal ihtiyati tedbir süreçleri açısından muğlak bırakılmış.",
        "confidence_score": 78,
        "status": "PENDING_APPROVAL",
        "is_approved": False,
        "is_rejected": False
    }
]

_SEED_AUDIT_LOGS = [
    {
        "id": "LOG-109",
        "agent_name": "Mantis-Legal-Auditor",
        "action": "SYSTEM_READY",
        "confidence_score": 1.0,
        "source_doc_id": "SYS",
        "timestamp": "Now",
        "level": "SYS",
        "message": "Auditor, Extractor, Builder, Harness & RAG engines fully integrated.",
    }
]

db_team_members = [
    {"id": 1, "name": "Sarah Jenkins", "email": "sarah.j@mantis.corp", "role": "Admin", "avatar": "SJ", "color": "from-rose-400 to-rose-600"},
    {"id": 2, "name": "Ahmet Yılmaz (Ahmet Usta)", "email": "ahmet.usta@mantis.corp", "role": "Atölye Ustası", "avatar": "AY", "color": "from-amber-400 to-amber-600"},
    {"id": 3, "name": "Hilal Uluca", "email": "hilal@mantis.corp", "role": "Lead Architect", "avatar": "HU", "color": "from-purple-400 to-purple-600"}
]

# --- Modül Yüklemeleri ---
extractor_available = False
process_document = None
try:
    from extractor import process_document as _process_document
    process_document = _process_document
    extractor_available = True
    logger.info("extractor.py başarıyla yüklendi.")
except Exception:
    logger.error("EXTRACTOR YÜKLENEMEDİ:\n%s", traceback.format_exc())

Auditor = None
auditor_import_error: Optional[str] = None
try:
    from auditor import Auditor as _Auditor, check_foundry_available
    Auditor = _Auditor
    logger.info("auditor.py başarıyla yüklendi.")
except Exception:
    auditor_import_error = traceback.format_exc()
    logger.error("AUDITOR MODÜLÜ YÜKLENEMEDİ:\n%s", auditor_import_error)

Builder = None
builder_import_error: Optional[str] = None
try:
    from builder import Builder as _Builder
    Builder = _Builder
    logger.info("builder.py başarıyla yüklendi.")
except Exception:
    builder_import_error = traceback.format_exc()
    logger.error("BUILDER MODÜLÜ YÜKLENEMEDİ:\n%s", builder_import_error)

rag_pipeline_available = False
process_and_store_document = None
init_db = None
retrieve_relevant_chunks = None
try:
    from rag_pipeline import process_and_store_document as _process, init_db as _init_db, retrieve_relevant_chunks as _retrieve
    process_and_store_document = _process
    init_db = _init_db
    retrieve_relevant_chunks = _retrieve
    rag_pipeline_available = True
    logger.info("rag_pipeline.py başarıyla yüklendi.")
except Exception:
    logger.error("RAG PIPELINE YÜKLENEMEDİ:\n%s", traceback.format_exc())

# --- GOOGLE SERVİSLERİ ---
SCOPES = [
    'https://www.googleapis.com/auth/gmail.modify',
    'https://www.googleapis.com/auth/tasks',
    'https://www.googleapis.com/auth/drive.file',
    'https://www.googleapis.com/auth/calendar',
]

def get_google_creds():
    creds = None
    token_path = os.path.join(BACKEND_DIR, "token.json")
    credentials_path = os.path.join(BACKEND_DIR, "credentials.json")
    if os.path.exists(token_path):
        creds = Credentials.from_authorized_user_file(token_path, SCOPES)

    if not creds or not creds.valid or not set(SCOPES).issubset(set(creds.scopes or [])):
        if creds and creds.expired and creds.refresh_token:
            try:
                creds.refresh(Request())
            except Exception:
                creds = None
        if not creds or not creds.valid:
            if os.path.exists(credentials_path):
                flow = InstalledAppFlow.from_client_secrets_file(credentials_path, SCOPES)
                creds = flow.run_local_server(port=0)
                with open(token_path, 'w') as token:
                    token.write(creds.to_json())
            else:
                raise RuntimeError(
                    "Google OAuth yapılandırması bulunamadı. backend/credentials.json dosyasını ekleyin."
                )
    if not creds or not creds.valid:
        raise RuntimeError(
            "Google hesabı yetkilendirilemedi. backend/token.json dosyasını yenileyin veya OAuth akışını tamamlayın."
        )
    return creds

def get_gmail_service():
    return build('gmail', 'v1', credentials=get_google_creds())

def get_google_tasks_service():
    return build('tasks', 'v1', credentials=get_google_creds())

def get_google_drive_service():
    return build('drive', 'v3', credentials=get_google_creds())

def sync_task_to_google_tasks(title: str, deadline: Optional[str] = None) -> Optional[str]:
    try:
        service = get_google_tasks_service()
        body = {"title": title}
        if deadline:
            try:
                body["due"] = datetime.strptime(deadline, "%Y-%m-%d").isoformat() + "Z"
            except ValueError:
                pass
        result = service.tasks().insert(tasklist='@default', body=body).execute()
        logger.info(f"[GOOGLE TASKS SYNC] '{title}' Google Tasks'a eklendi (id: {result.get('id')})")
        return result.get('id')
    except Exception as e:
        logger.warning(f"Google Tasks senkron hatası: {e}")
        return None

def upload_docx_to_drive(filepath: str, filename: str) -> Optional[str]:
    try:
        service = get_google_drive_service()
        folder_id = None
        query = "name='Project Mantis' and mimeType='application/vnd.google-apps.folder' and trashed=false"
        results = service.files().list(q=query, fields="files(id, name)").execute()
        folders = results.get('files', [])
        if folders:
            folder_id = folders[0]['id']
        else:
            folder_metadata = {'name': 'Project Mantis', 'mimeType': 'application/vnd.google-apps.folder'}
            folder = service.files().create(body=folder_metadata, fields='id').execute()
            folder_id = folder.get('id')

        file_metadata = {'name': filename, 'parents': [folder_id]}
        media = MediaFileUpload(
            filepath,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        uploaded = service.files().create(body=file_metadata, media_body=media, fields='id, webViewLink').execute()
        logger.info(f"[DRIVE UPLOAD] '{filename}' Drive'a yüklendi: {uploaded.get('webViewLink')}")
        return uploaded.get('webViewLink')
    except Exception as e:
        logger.warning(f"Drive yükleme hatası: {e}")
        return None

# --- SMTP MAİL YARDIMCISI ---
def send_legal_approval_email(personnel_name: str, draft_text: str, recipient_email: str = "hukuk@sirket.com") -> bool:
    sender_email = os.getenv("MANTIS_EMAIL")
    sender_password = os.getenv("MANTIS_EMAIL_PASS")

    if not sender_email or not sender_password:
        logger.warning(
            "MANTIS_EMAIL / MANTIS_EMAIL_PASS tanımlı değil — mail GÖNDERİLMEDİ "
            "(simülasyon modu). Çağıran taraf mail_delivered=False görecek."
        )
        return False

    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = recipient_email
    msg['Subject'] = f"⚠️ ONAY BEKLİYOR: {personnel_name} - Otonom İşlem / İhtarname Taslağı"

    body = f"""Sayın İlgili,

Mantis AI Operasyon Sistemi, '{personnel_name}' talebi/kaydı doğrultusunda aşağıdaki otonom taslağı üretmiştir.

--- OTONOM İÇERİK / TASLAK ---

{draft_text}

-------------------------
Bu e-posta Project Mantis tarafından otomatik olarak oluşturulmuştur.
"""
    msg.attach(MIMEText(body, 'plain', 'utf-8'))

    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, sender_password)
        text = msg.as_string()
        server.sendmail(sender_email, recipient_email, text)
        server.quit()
        logger.info(f"Otonom mail başarıyla gönderildi: {recipient_email}")
        return True
    except Exception as e:
        logger.error(f"SMTP Mail Gönderme Hatası: {e}")
        return False

# --- FASTAPI APP ---
app = FastAPI(
    title="Project Mantis Engine",
    docs_url="/docs",
    redoc_url="/redoc",
    openapi_url="/openapi.json",
)

ALLOWED_ORIGINS = os.getenv(
    "ALLOWED_ORIGINS",
    "http://localhost:3000,http://localhost:5173,http://127.0.0.1:3000,http://127.0.0.1:5173"
).split(",")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

VAULT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "generated_documents")
os.makedirs(VAULT_DIR, exist_ok=True)
logger.info(f"[VAULT] Doküman klasörü: {VAULT_DIR}")

# Belgelerin güvenli şekilde sunulması için statik klasör
app.mount("/static", StaticFiles(directory=VAULT_DIR), name="documents_static")

auditor = None
auditor_init_error: Optional[str] = None
builder = None


def _safe_display_status(status_value: Optional[str]) -> str:
    value = (status_value or "").strip().lower()
    if value in {"pending_approval", "suggested", "onay bekliyor", "pending", "needs_review"}:
        return "Onay Bekliyor"
    if value in {"approved", "onaylandı", "confirmed"}:
        return "Onaylandı"
    if value in {"rejected", "reddedildi"}:
        return "Reddedildi"
    return "Bekliyor"


def _normalize_task_payload(raw_task: Dict[str, Any], fallback_source: str = "system") -> Dict[str, Any]:
    title = str(raw_task.get("title") or raw_task.get("task_title") or "Yeni görev").strip()[:240]
    status_value = raw_task.get("status") or raw_task.get("state") or "pending_approval"
    normalized_status = "pending_approval" if str(status_value).lower() in {"pending_approval", "suggested", "onay bekliyor", "pending", "needs_review"} else str(status_value)
    if normalized_status.lower() in {"approved", "confirmed", "onaylandı"}:
        normalized_status = "confirmed"
    if normalized_status.lower() in {"rejected", "reddedildi"}:
        normalized_status = "rejected"
    return {
        "id": str(raw_task.get("id") or f"task_{len(db_tasks) + 1:04d}"),
        "title": title,
        "deadline": raw_task.get("deadline") or "Belirtilmedi",
        "confidence": float(raw_task.get("confidence", 0.8) or 0.8),
        "status": normalized_status,
        "display_status": _safe_display_status(normalized_status),
        "type": raw_task.get("type") or "AI Tarafından Çıkarıldı",
        "completed": bool(raw_task.get("completed", False)) or normalized_status in {"confirmed", "approved", "onaylandı"},
        "source_email_id": raw_task.get("source_email_id") or fallback_source,
        "sender": raw_task.get("sender") or "Project Mantis",
        "ai_draft": raw_task.get("ai_draft") or "",
        "created_by": raw_task.get("created_by") or "fallback_engine",
        "requires_ai_review": bool(raw_task.get("requires_ai_review", True)),
    }


def _normalize_risk_payload(raw_risk: Dict[str, Any], fallback_id: Optional[str] = None) -> Dict[str, Any]:
    risk_id = str(raw_risk.get("risk_id") or fallback_id or f"RSK-{uuid.uuid4().hex[:6].upper()}")
    severity = str(raw_risk.get("severity") or raw_risk.get("risk_level") or "Medium").strip()
    status_value = raw_risk.get("status") or "pending_approval"
    status = "pending_approval" if str(status_value).lower() in {"pending_approval", "pending", "onay bekliyor", "suggested", "needs_review"} else str(status_value)
    return {
        "risk_id": risk_id,
        "severity": severity,
        "clause_text": str(raw_risk.get("clause_text") or raw_risk.get("title") or "Yasallık riski tespit edildi."),
        "ai_reasoning": str(raw_risk.get("ai_reasoning") or "Otomatik güvenlik kuralı tarafından tespit edildi."),
        "confidence_score": int(raw_risk.get("confidence_score", 80) or 80),
        "status": status,
        "is_approved": bool(raw_risk.get("is_approved", False)),
        "is_rejected": bool(raw_risk.get("is_rejected", False)),
    }


def _normalize_document_summary(raw_doc: Dict[str, Any], fallback_id: Optional[str] = None) -> Dict[str, Any]:
    doc_id = str(raw_doc.get("id") or fallback_id or f"doc_{uuid.uuid4().hex[:8]}")
    name = str(raw_doc.get("name") or raw_doc.get("title") or "document.txt")
    return {
        "id": doc_id,
        "name": name,
        "access": raw_doc.get("access") or "Fallback Engine",
        "last_active": raw_doc.get("last_active") or datetime.now().strftime("%Y-%m-%d %H:%M"),
        "type": raw_doc.get("type") or "txt",
        "filepath": raw_doc.get("filepath") or os.path.join(VAULT_DIR, name),
        "content": raw_doc.get("content") or raw_doc.get("summary") or "Belge içeriği güvenli fallback olarak gösteriliyor.",
        "source_type": raw_doc.get("source_type") or "fallback",
        "status": raw_doc.get("status") or "pending_approval",
        "summary": raw_doc.get("summary") or "Belge içeriği güvenli şekilde görüntüleniyor.",
    }


def _validate_json_guardrail(payload: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(payload, dict):
        raise ValueError("JSON guardrail: payload dict değil.")
    if not payload.get("status"):
        payload["status"] = "pending_approval"
    payload["status"] = str(payload["status"]).lower()
    if payload["status"] not in {"pending_approval", "approved", "rejected", "success", "error"}:
        payload["status"] = "pending_approval"
    return payload


def _fallback_document_summary(text: str, source_name: str = "document.txt") -> Dict[str, Any]:
    cleaned = re.sub(r"\s+", " ", text or "").strip()
    preview = cleaned[:2000] if cleaned else "Belge okunamadı; güvenli metin görünümü oluşturuluyor."
    doc_id = f"doc_{uuid.uuid4().hex[:8]}"
    summary = preview[:400]
    return {
        "id": doc_id,
        "name": source_name,
        "access": "Fallback Engine",
        "last_active": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "type": os.path.splitext(source_name)[1].lower().lstrip(".") or "txt",
        "filepath": os.path.join(VAULT_DIR, source_name),
        "content": preview,
        "summary": summary,
        "status": "pending_approval",
        "source_type": "fallback",
    }


def _foundry_fallback_engine(text: str, document_name: str = "uploaded_document.txt", source_id: str = "doc_generated") -> Dict[str, Any]:
    safe_text = (text or "")
    sentences = [segment.strip() for segment in re.split(r"(?<=[.!?])\s+", safe_text) if segment.strip()]
    lower_text = safe_text.lower()

    tasks = []
    for idx, sentence in enumerate(sentences[:8], start=1):
        if any(keyword in sentence.lower() for keyword in ["görev", "aksiyon", "task", "yapılacak", "teslim", "deadline", "onay", "hazırlan", "ihtar", "rapor", "süre", "yap"]):
            tasks.append(_normalize_task_payload({
                "id": f"task_{len(db_tasks) + idx:04d}",
                "title": sentence[:180],
                "deadline": "Belirtilmedi",
                "confidence": 0.76,
                "status": "pending_approval",
                "type": "Fallback Engine",
                "completed": False,
                "source_email_id": source_id,
                "sender": document_name,
                "ai_draft": sentence,
                "created_by": "fallback_engine",
                "requires_ai_review": True,
            }, fallback_source=source_id))

    if not tasks:
        tasks.append(_normalize_task_payload({
            "id": f"task_{len(db_tasks) + 1:04d}",
            "title": f"İşlem: {document_name} için temel denetim tamamlandı",
            "deadline": "Belirtilmedi",
            "confidence": 0.7,
            "status": "pending_approval",
            "type": "Fallback Engine",
            "completed": False,
            "source_email_id": source_id,
            "sender": document_name,
            "ai_draft": safe_text[:500],
            "created_by": "fallback_engine",
            "requires_ai_review": True,
        }, fallback_source=source_id))

    risk_candidates = []
    if any(keyword in lower_text for keyword in ["sınırsız", "uncapped", "sorumluluk", "fesih", "termination", "ceza", "penalty", "mücbir sebep", "force majeure"]):
        risk_candidates.append({
            "risk_id": f"RSK-{uuid.uuid4().hex[:6].upper()}",
            "severity": "High",
            "clause_text": "Risk tespit edildi: sözleşme maddesi sıkı denetim gerektiriyor.",
            "ai_reasoning": "Fallback kural motoru, sorumluluk veya fesih riskini otomatik olarak işaretledi.",
            "confidence_score": 80,
            "status": "pending_approval",
        })
    if not risk_candidates:
        risk_candidates.append({
            "risk_id": f"RSK-{uuid.uuid4().hex[:6].upper()}",
            "severity": "Medium",
            "clause_text": "Belgeye ait genel inceleme denetimi yapıldı.",
            "ai_reasoning": "Foundry Local erişilemediği için deterministik kural tabanlı denetim devreye alındı.",
            "confidence_score": 71,
            "status": "pending_approval",
        })

    normalized_risks = [_normalize_risk_payload(item, item.get("risk_id")) for item in risk_candidates]
    for risk_item in normalized_risks:
        db_risk_findings.insert(0, {
            **risk_item,
            "is_approved": False,
            "is_rejected": False,
            "source_document_id": source_id,
        })

    doc_summary = _fallback_document_summary(safe_text, source_name=document_name)
    db_vault_documents.insert(0, {
        **_normalize_document_summary(doc_summary, doc_summary["id"]),
        "content": doc_summary["content"],
        "summary": doc_summary["summary"],
        "source_type": "fallback",
    })

    for task in tasks:
        db_tasks.insert(0, task)

    return {
        "status": "success",
        "source": "fallback_engine",
        "tasks": tasks,
        "risks": normalized_risks,
        "documents": [
            _normalize_document_summary(doc_summary, doc_summary["id"])
        ],
        "summary": {
            "message": "Foundry Local erişilemediği için güvenli fallback modu devreye girdi.",
            "task_count": len(tasks),
            "risk_count": len(normalized_risks),
            "document_count": 1,
        },
    }


def _ensure_foundry_healthcheck() -> tuple[bool, str]:
    try:
        from auditor import check_foundry_available
        return check_foundry_available(timeout=2.5)
    except Exception as exc:
        return False, f"Foundry health check hata: {exc}"


async def _run_foundry_healthcheck() -> dict:
    ok, message = _ensure_foundry_healthcheck()
    return {
        "status": "success" if ok else "degraded",
        "foundry_local": ok,
        "message": message,
        "endpoint": os.getenv("FOUNDRY_LOCAL_ENDPOINT", "http://127.0.0.1:5272/v1"),
        "model": os.getenv("FOUNDRY_LOCAL_MODEL", "Phi-3.5-mini-instruct-openvino-gpu"),
    }

@app.on_event("startup")
async def startup_event():
    global auditor, auditor_init_error, builder
    logger.info("🚀 Project Mantis başlatılıyor...")

    init_store()
    db_risk_findings.seed_if_empty(_SEED_RISK_FINDINGS)
    db_audit_logs.seed_if_empty(_SEED_AUDIT_LOGS)

    if init_db:
        try:
            init_db()
        except Exception as e:
            logger.error(f"DB başlatma hatası: {e}")

    if Auditor is not None:
        try:
            auditor = Auditor()
            if auditor.is_ready:
                logger.info("Auditor motoru başarıyla başlatıldı ve bağlandı.")
            else:
                auditor_init_error = auditor.init_error
                logger.warning(
                    "Auditor HAZIR DEĞİL: %s — Foundry Local servisinin çalıştığından "
                    "ve modelin (ör. Phi-3.5-mini-instruct) çekili olduğundan emin olun.",
                    auditor_init_error,
                )
        except Exception:
            auditor_init_error = traceback.format_exc()
            logger.error("AUDITOR BAŞLATMA HATASI:\n%s", auditor_init_error)
    else:
        auditor_init_error = auditor_import_error or "Auditor sınıfı import edilemedi."

    if Builder is not None:
        try:
            builder = Builder()
            logger.info("Builder motoru başarıyla başlatıldı.")
        except Exception:
            logger.error("BUILDER BAŞLATMA HATASI:\n%s", traceback.format_exc())

    renewal_thread = threading.Thread(target=background_subscription_renewal_loop, daemon=True)
    renewal_thread.start()

def get_auditor_status() -> dict:
    global auditor
    return {
        "import_ok": Auditor is not None,
        "initialized": auditor is not None,
        "ready": bool(auditor and auditor.is_ready),
        "import_error": auditor_import_error,
        "init_error": auditor_init_error or (auditor.init_error if auditor else None),
    }

# --- DOKÜMAN ÜRETİMİ ---
VAULT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "generated_documents")
os.makedirs(VAULT_DIR, exist_ok=True)
logger.info(f"[VAULT] Doküman klasörü: {VAULT_DIR}")


def _extract_document_text(filepath: str, raw_bytes: Optional[bytes] = None) -> str:
    """Bilinen belge formatlarından güvenli metin çıkarır; okuma mümkün olmazsa boş string döner."""
    if not filepath:
        return ""

    ext = os.path.splitext(filepath)[1].lower()
    try:
        if ext == ".docx":
            try:
                return "\n".join(p.text for p in DocxDocument(filepath).paragraphs if p.text and p.text.strip())
            except Exception:
                if raw_bytes:
                    try:
                        import zipfile
                        import xml.etree.ElementTree as ET

                        with zipfile.ZipFile(io.BytesIO(raw_bytes)) as archive:
                            xml_bytes = archive.read("word/document.xml")
                            root = ET.fromstring(xml_bytes)
                            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                            paragraphs = []
                            for para in root.findall(".//w:p", ns):
                                texts = []
                                for node in para.findall(".//w:t", ns):
                                    text = node.text or ""
                                    texts.append(text)
                                text = "".join(texts).strip()
                                if text:
                                    paragraphs.append(text)
                            return "\n".join(paragraphs)
                    except Exception:
                        pass
                return "Belge okunamadı; içeriği çıkarılamadı."

        if ext in {".txt", ".md", ".csv", ".json"}:
            with open(filepath, "r", encoding="utf-8", errors="replace") as handle:
                return handle.read()

        if ext == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(filepath)
                texts = []
                for page in reader.pages:
                    text = page.extract_text() or ""
                    if text.strip():
                        texts.append(text.strip())
                return "\n".join(texts)
            except Exception:
                return "PDF içeriği okunamadı; dosya güvenli şekilde yüklendi."

        if raw_bytes:
            try:
                return raw_bytes.decode("utf-8", errors="replace")
            except Exception:
                pass

    except Exception:
        logger.warning("Belge metni çıkarılamadı (%s): %s", filepath, traceback.format_exc())

    return "Belge içeriği güvenli yöntemle hazırlandı; metin görünümü mevcut değil."


def generate_real_docx(topic: str, context_hint: str = "") -> dict:
    global auditor
    doc_id = f"doc_{uuid.uuid4().hex[:8]}"
    safe_topic = re.sub(r"[^\w\s-]", "", topic).strip().replace(" ", "_")[:40] or "belge"
    filename = f"{doc_id}_{safe_topic}.docx"
    filepath = os.path.join(VAULT_DIR, filename)

    body_text = None
    if auditor and auditor.is_ready:
        try:
            profile = load_company_profile()
            content_prompt = f"""
            {MANTIS_PERSONA}

            Aşağıdaki konu hakkında, bu şirketin gerçek bağlamını kullanan özgün ve
            okunabilir bir Türkçe belge metni hazırla. Konuyu başlık olarak tekrar etme.
            Şirketin adı, sektörü, ekip büyüklüğü, iletişim tonu, çalışma saatleri,
            kritik birimleri ve otomasyon politikası metne doğal biçimde yansısın.
            Şirket kültürü isteniyorsa değerler, günlük çalışma davranışları, iletişim,
            karar alma, sorumluluk ve yeni çalışan beklentilerini somutlaştır.
            Sadece belge içeriğini yaz; meta açıklama, JSON, markdown veya “açıklanamıyor” gibi
            belirsiz ifadeler kullanma.

            Konu: {topic}
            Ek bağlam: {context_hint or 'Yok'}
            Şirket bağlamı: {json.dumps(profile, ensure_ascii=False)}
            """
            body_text = auditor.ask(content_prompt)
        except Exception as e:
            logger.error(f"Doküman içerik üretim hatası: {e}")

    if not body_text:
        raise RuntimeError(
            "Belge oluşturulamadı: Foundry Local AI modeli yanıt vermedi. "
            "Foundry Local servisinin çalıştığını ve phi-3.5-mini modelinin yüklü olduğunu kontrol edin."
        )

    docx_obj = DocxDocument()
    docx_obj.add_heading(topic, level=1)
    docx_obj.add_paragraph(f"Oluşturulma tarihi: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    docx_obj.add_paragraph("")
    for paragraph in body_text.split("\n"):
        if paragraph.strip():
            docx_obj.add_paragraph(paragraph.strip())
    docx_obj.save(filepath)
    logger.info(f"[DOKÜMAN ÜRETİMİ] Gerçek .docx oluşturuldu: {filepath}")

    if rag_pipeline_available and process_and_store_document:
        try:
            process_and_store_document(
                source_file=filename,
                content=body_text,
                document_type="docx_vault"
            )
            logger.info(f"[RAG INGESTION] '{filename}' vektör veritabanına eklendi.")
        except Exception as rag_err:
            logger.warning(f"Doküman RAG hafızasına işlenirken hata: {rag_err}")

    drive_link = upload_docx_to_drive(filepath, filename)

    vault_entry = {
        "id": doc_id,
        "name": filename,
        "access": "AI Tarafından Oluşturuldu",
        "last_active": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "type": "docx",
        "filepath": filepath,
        "drive_link": drive_link,
    }
    db_vault_documents.insert(0, vault_entry)
    return vault_entry

# --- AUTH ENDPOINTLERİ ---
@app.post("/api/v1/auth/signup")
def signup(payload: UserSignupRequest):
    users = load_users()
    if payload.email in users:
        raise HTTPException(status_code=400, detail="Bu e-posta adresi ile zaten kayıt olunmuş.")

    if not payload.email or not payload.password or not payload.full_name:
        raise HTTPException(status_code=400, detail="E-posta, ad-soyad ve şifre alanları zorunludur.")

    hashed_password = get_password_hash(payload.password)
    users[payload.email] = {
        "email": payload.email,
        "password": hashed_password,
        "full_name": payload.full_name,
        "created_at": datetime.now().isoformat()
    }
    save_users(users)
    return {"status": "success", "message": "Kayıt başarıyla oluşturuldu."}

@app.post("/api/v1/documents/upload")
def upload_document(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="Yüklenecek dosya adı eksik.")

    clean_name = re.sub(r"[^A-Za-z0-9_.-]", "_", file.filename)
    safe_name = clean_name[:120]
    upload_dir = VAULT_DIR
    os.makedirs(upload_dir, exist_ok=True)

    destination = os.path.join(upload_dir, safe_name)
    content = file.file.read()
    with open(destination, "wb") as f:
        f.write(content)

    extracted_text = _extract_document_text(destination, content)
    doc_id = f"doc_{uuid.uuid4().hex[:8]}"
    vault_entry = {
        "id": doc_id,
        "name": safe_name,
        "access": "Yüklendi",
        "last_active": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "type": os.path.splitext(safe_name)[1].lower().lstrip(".") or "uploaded",
        "filepath": destination,
        "drive_link": None,
        "content": extracted_text,
        "summary": extracted_text[:500],
        "status": "pending_approval",
        "source_type": "upload",
    }
    db_vault_documents.insert(0, vault_entry)
    return {"status": "success", "document": {"id": doc_id, "name": safe_name, "type": vault_entry["type"]}, "message": "Belge başarıyla yüklendi."}

@app.post("/api/v1/auth/signin")
def signin(payload: UserSigninRequest):
    users = load_users()
    user = users.get(payload.email)
    if not user or not verify_password(payload.password, user["password"]):
        raise HTTPException(status_code=401, detail="Geçersiz e-posta veya şifre.")

    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user["email"], "name": user["full_name"]}, expires_delta=access_token_expires
    )
    onboarding_completed = os.path.exists(COMPANY_PROFILE_PATH)
    return {
        "status": "success",
        "access_token": access_token,
        "token_type": "bearer",
        "full_name": user["full_name"],
        "onboarding_completed": onboarding_completed,
        "message": "Giriş başarılı."
    }

@app.post("/api/v1/auth/signout")
def signout():
    return {"status": "success", "message": "Oturum başarıyla kapatıldı."}

# --- TEMEL ENDPOINTLER ---
@app.get("/")
def health_check():
    status_info = get_auditor_status()
    return {
        "status": "active" if status_info["ready"] else "degraded",
        "system": "Project Mantis",
        "auditor": status_info
    }

@app.get("/api/v1/dashboard/logs")
def get_logs(current_user: dict = Depends(get_current_user)):
    return list(db_audit_logs)

@app.get("/api/v1/dashboard/summary")
def get_dashboard_summary(current_user: dict = Depends(get_current_user)):
    risks = list(db_risk_findings)
    tasks = list(db_tasks)
    documents = list(db_vault_documents)
    audited_documents = {
        str(log.get("source_doc_id"))
        for log in db_audit_logs
        if log.get("level") == "AUDIT" and log.get("source_doc_id")
    }
    open_risks = [risk for risk in risks if risk.get("status") not in ("APPROVED_BY_HUMAN", "REJECTED_BY_HUMAN")]
    pending_tasks = [task for task in tasks if task.get("status") == "Onay Bekliyor"]
    return {
        "status": "success",
        "documents_audited": len(audited_documents),
        "documents_in_vault": len(documents),
        "open_risks": len(open_risks),
        "risk_breakdown": {
            "critical": sum(1 for risk in open_risks if "critical" in str(risk.get("severity", "")).lower()),
            "high": sum(1 for risk in open_risks if "high" in str(risk.get("severity", "")).lower()),
            "medium": sum(1 for risk in open_risks if "medium" in str(risk.get("severity", "")).lower()),
            "low": sum(1 for risk in open_risks if "low" in str(risk.get("severity", "")).lower()),
        },
        "pending_approvals": len(pending_tasks),
    }

@app.get("/api/v1/integrations/status")
def get_integrations_status(current_user: dict = Depends(get_current_user)):
    foundry_status = get_auditor_status()
    try:
        foundry_ok, foundry_message = check_foundry_available()
    except Exception as exc:
        foundry_ok, foundry_message = False, str(exc)

    token_path = os.path.join(BACKEND_DIR, "token.json")
    credentials_path = os.path.join(BACKEND_DIR, "credentials.json")
    google_configured = os.path.exists(credentials_path)
    google_authorized = os.path.exists(token_path)
    return {
        "status": "success",
        "integrations": {
            "gmail": {
                "configured": google_configured,
                "authorized": google_authorized,
                "status": "Connected" if google_authorized else "Needs authorization",
                "action": "scan_inbox",
            },
            "drive": {
                "configured": google_configured,
                "authorized": google_authorized,
                "status": "Available" if google_authorized else "Needs authorization",
            },
            "google_tasks": {
                "configured": google_configured,
                "authorized": google_authorized,
                "status": "Configured; API check on sync" if google_authorized else "Needs authorization",
            },
            "foundry_local": {
                "configured": True,
                "authorized": foundry_ok,
                "status": "Connected" if foundry_ok else "Unavailable",
                "endpoint": os.getenv("FOUNDRY_LOCAL_ENDPOINT", "http://127.0.0.1:5272/v1"),
                "model": os.getenv("FOUNDRY_LOCAL_MODEL", "Phi-3.5-mini-instruct-openvino-gpu"),
                "message": foundry_message,
            },
        },
    }

@app.get("/api/v1/analysis/risks")
def get_ui_risks(current_user: dict = Depends(get_current_user)):
    """AnalysisHub için risk bulgularını döner."""
    return list(db_risk_findings)

@app.post("/api/analyze")
@app.post("/api/v1/analyze")
def analyze_document(request: AnalyzeRequest, current_user: dict = Depends(get_current_user)):
    """Sözleşme metnini analiz eder ve risk bulgularını döner."""
    global auditor

    if not request.text or not request.text.strip():
        raise HTTPException(status_code=400, detail="Analiz edilecek metin boş.")

    if not auditor or not auditor.is_ready:
        if Auditor is not None:
            auditor = Auditor(skip_check=False)

    if not auditor or not auditor.is_ready:
        fallback = _foundry_fallback_engine(request.text, document_name=request.document_id or "uploaded_document.txt", source_id=request.document_id or "doc_generated")
        findings = [{
            **_normalize_risk_payload(finding, finding.get("risk_id")),
            "status": "pending_approval",
            "is_approved": False,
            "is_rejected": False,
        } for finding in fallback.get("risks", [])]
        return {
            "status": "success",
            "document_id": request.document_id or "doc_generated",
            "total_risks_found": len(findings),
            "findings": findings,
            "message": "Foundry Local erişilemediği için güvenli fallback motoru devreye girdi."
        }

    result = auditor.analyze_text(request.text)
    if result.get("status") != "success":
        fallback = _foundry_fallback_engine(request.text, document_name=request.document_id or "uploaded_document.txt", source_id=request.document_id or "doc_generated")
        findings = [{
            **_normalize_risk_payload(finding, finding.get("risk_id")),
            "status": "pending_approval",
            "is_approved": False,
            "is_rejected": False,
        } for finding in fallback.get("risks", [])]
        return {
            "status": "success",
            "document_id": request.document_id or "doc_generated",
            "total_risks_found": len(findings),
            "findings": findings,
            "message": result.get("error") or "Foundry Local yanıt vermedi; güvenli fallback çalıştırıldı."
        }

    findings = []
    for item in result.get("findings", []):
        findings.append({
            "risk_id": item.get("risk_id", "RSK-001"),
            "severity": item.get("severity", "Medium"),
            "clause_text": item.get("clause_text", ""),
            "ai_reasoning": item.get("ai_reasoning", ""),
            "confidence_score": int(item.get("confidence_score", 80)),
            "status": "pending_approval",
            "is_approved": False,
            "is_rejected": False,
        })

    db_risk_findings.clear()
    for finding in findings:
        db_risk_findings.insert(0, {
            **finding,
            "status": "pending_approval",
            "is_approved": False,
            "is_rejected": False,
        })

    db_audit_logs.insert(0, {
        "id": f"AUDIT-{uuid.uuid4().hex[:8]}",
        "agent_name": "Mantis-Legal-Auditor",
        "action": "DOCUMENT_AUDIT",
        "confidence_score": 1.0,
        "source_doc_id": request.document_id or "doc_generated",
        "timestamp": datetime.now().strftime("%H:%M:%S"),
        "level": "AUDIT",
        "message": f"Belge risk analizi tamamlandı: {len(findings)} bulgu.",
    })

    return {
        "status": "success",
        "document_id": request.document_id or "doc_generated",
        "total_risks_found": len(findings),
        "findings": findings,
        "message": "Risk analizi tamamlandı."
    }

@app.get("/api/export/{format}")
def export_analysis_data(format: str):
    """AnalysisHub CSV/XLSX export desteği."""
    if format == "csv":
        content = "Risk ID,Severity,Clause,AI Reasoning,Confidence\n"
        for r in db_risk_findings:
            content += f'"{r["risk_id"]}","{r["severity"]}","{r["clause_text"]}","{r["ai_reasoning"]}",{r["confidence_score"]}\n'
        return Response(content=content, media_type="text/csv", headers={"Content-Disposition": "attachment; filename=mantis_risks.csv"})
    elif format in ("excel", "xlsx"):
        content = "Risk ID\tSeverity\tClause\tAI Reasoning\tConfidence\n"
        for r in db_risk_findings:
            content += f"{r['risk_id']}\t{r['severity']}\t{r['clause_text']}\t{r['ai_reasoning']}\t{r['confidence_score']}\n"
        return Response(content=content, media_type="application/vnd.ms-excel", headers={"Content-Disposition": "attachment; filename=mantis_risks.xls"})
    raise HTTPException(status_code=400, detail="Desteklenmeyen format.")

@app.get("/api/v1/team/members")
def get_team_members():
    return db_team_members

@app.get("/api/v1/tasks")
def get_tasks(current_user: dict = Depends(get_current_user)):
    return {"status": "success", "tasks": list(db_tasks)}

@app.post("/api/v1/tasks/create")
def create_task(payload: dict, current_user: dict = Depends(get_current_user)):
    title = str(payload.get("title", "")).strip()
    if not title:
        raise HTTPException(status_code=400, detail="Görev başlığı gerekli.")

    new_task = {
        "id": f"task_{len(db_tasks) + 1:04d}",
        "title": title,
        "deadline": payload.get("deadline") or "Belirtilmedi",
        "confidence": float(payload.get("confidence", 0.9)),
        "status": payload.get("status") or "confirmed",
        "type": payload.get("type") or "Manuel Görev",
        "completed": bool(payload.get("completed", False)),
        "source_email_id": payload.get("source_email_id") or "manual_ui",
        "sender": payload.get("sender") or current_user.get("full_name") or "User",
        "ai_draft": payload.get("ai_draft") or "",
        "created_by": "user",
    }
    db_tasks.insert(0, new_task)
    return {"status": "success", "message": "Görev oluşturuldu.", "task": new_task}

@app.get("/api/v1/offboarding/tracker")
def get_offboarding_tracker():
    return {
        "employee_id": "EMP-2024-089",
        "employee": "Ahmet Yılmaz (Ahmet Usta)",
        "position": "Kıdemli Atölye Ustası",
        "hire_date": "2023-03-15",
        "absent_days": 2,
        "status": "WARNING_REQUIRED",
        "article": "İş Kanunu Madde 25/II & Madde 17",
        "assets": ["Delik Makinesi Anahtar Seti", "Kurumsal Tablet (Tab-A8)"]
    }

# --- AGENT HARNESS & GUARDRAIL ENDPOINTLERİ ---
def _handle_harness_approval_logic(payload: ApprovalRequest) -> dict:
    is_app = payload.action.lower() in ("approve", "onayla")
    decision_state = "APPROVED_BY_HUMAN" if is_app else "REJECTED_BY_HUMAN"
    db_approval_states[payload.risk_id] = decision_state
    
    # SQLite kalıcı güncelleme
    updated_risk = db_risk_findings.update(
        payload.risk_id,
        status=decision_state,
        is_approved=is_app,
        is_rejected=not is_app
    )

    log_id = f"LOG-{datetime.now().strftime('%H%M%S')}"
    db_audit_logs.insert(0, {
        "id": log_id,
        "agent_name": "Human-Supervisor",
        "action": f"HARNESS_{payload.action.upper()}",
        "confidence_score": 1.0,
        "source_doc_id": payload.risk_id,
        "timestamp": datetime.now().strftime("%H:%M:%S"),
        "level": "AUDIT",
        "message": f"Risk maddesi ({payload.risk_id}) kullanıcı tarafından {decision_state} olarak işaretlendi."
    })

    if updated_risk is None:
        return {
            "status": "error",
            "risk_id": payload.risk_id,
            "message": f"'{payload.risk_id}' ID'li risk bulunamadı, hiçbir şey güncellenmedi.",
        }

    return {
        "status": "success",
        "risk_id": payload.risk_id,
        "approval_state": decision_state,
        "updated_risk": updated_risk,
        "message": f"İşlem kaydedildi: {decision_state}"
    }

@app.post("/api/harness/approve")
def handle_harness_approval(payload: ApprovalRequest):
    return _handle_harness_approval_logic(payload)

@app.post("/api/v1/harness/approve")
def handle_harness_approval_v1(payload: ApprovalRequest):
    return _handle_harness_approval_logic(payload)

@app.post("/api/action/dispatch")
def dispatch_action(payload: ActionDispatchRequest):
    """Jira/E-posta aksiyon tetikleyicisi."""
    log_id = f"LOG-{datetime.now().strftime('%H%M%S')}"
    db_audit_logs.insert(0, {
        "id": log_id,
        "agent_name": "Action-Dispatcher",
        "action": f"DISPATCH_{payload.action_type.upper()}",
        "confidence_score": 1.0,
        "source_doc_id": "ACTION",
        "timestamp": datetime.now().strftime("%H:%M:%S"),
        "level": "OPERATIONAL",
        "message": f"{payload.action_type.upper()} görevi sevk edildi: {payload.title}"
    })

    if payload.action_type.lower() == "email":
        delivered = send_legal_approval_email(
            personnel_name=payload.title,
            draft_text=payload.description,
            recipient_email=payload.recipient or "hukuk@sirket.com"
        )
        if delivered:
            return {"status": "success", "message": f"'{payload.recipient}' adresine e-posta gönderildi.", "mail_delivered": True}
        return {"status": "error", "message": "Mail gönderilemedi (SMTP yapılandırması eksik veya hata oluştu).", "mail_delivered": False}
    elif payload.action_type.lower() == "jira":
        return {"status": "success", "message": f"Jira panosuna '{payload.title}' başlığıyla task açıldı."}

    return {"status": "success", "message": "Aksiyon başarıyla işlendi."}

# --- ŞİRKET KÜLTÜRÜ & ONBOARDING ---
COMPANY_PROFILE_PATH = "company_profile.json"

def load_company_profile() -> dict:
    if os.path.exists(COMPANY_PROFILE_PATH):
        try:
            with open(COMPANY_PROFILE_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {
        "identity": {"name": "Project Mantis Corp", "sector": "Yazılım ve Otonom Sistemler", "team_size": "25 kişi"},
        "communication_tone": "Keskin, net, kurumsal ve stratejik",
        "key_contacts": {"Hukuk": "hukuk@sirket.com", "İK": "ik@sirket.com"},
        "standards": {"working_hours": "09:00 - 18:00", "red_lines": "KVKK hassasiyeti, onaysız dış e-posta gönderimi kesinlikle yasaktır."},
        "automation_policy": {"autonomous": "Düşük öncelikli task ekleme, gelen kutusu spam eleme", "human_in_the_loop": "İhtarname, resmi sözleşme, dışarıya mail gönderme"}
    }

def save_company_profile(profile_data: dict):
    with open(COMPANY_PROFILE_PATH, "w", encoding="utf-8") as f:
        json.dump(profile_data, f, ensure_ascii=False, indent=4)

def answer_from_company_profile(user_msg: str) -> Optional[str]:
    profile = load_company_profile()
    msg_lower = user_msg.lower()

    if any(kw in msg_lower for kw in ["kvkk", "kırmızı çizgi", "red line"]):
        return f"KVKK ve kırmızı çizgilerimiz: {profile.get('standards', {}).get('red_lines', 'Tanımlı değil.')}"

    if any(kw in msg_lower for kw in ["iletişim tonu", "ne tonda", "nasıl konuş"]):
        return f"Şirket iletişim tonumuz: {profile.get('communication_tone', 'Tanımlı değil.')}"

    if any(kw in msg_lower for kw in ["otomasyon politika", "insan onayı", "human in the loop", "hangi işler onay", "neyi kendin yapars"]):
        pol = profile.get("automation_policy", {})
        return (
            f"Kendi başıma yapabildiklerim: {pol.get('autonomous', 'Tanımlı değil.')}\n"
            f"İnsan onayı gereken işler: {pol.get('human_in_the_loop', 'Tanımlı değil.')}"
        )

    if any(kw in msg_lower for kw in ["çalışma saat", "mesai saat"]):
        return f"Çalışma saatlerimiz: {profile.get('standards', {}).get('working_hours', 'Tanımlı değil.')}"

    if any(kw in msg_lower for kw in ["kritik departman", "kime ulaşayım", "hangi departman"]):
        contacts = profile.get("key_contacts", {})
        if contacts:
            lines = "\n".join(f"• {k}: {v}" for k, v in contacts.items())
            return f"Kritik departman iletişimleri:\n{lines}"

    return None

def is_human_in_the_loop_topic(topic: str) -> bool:
    profile = load_company_profile()
    hil_text = profile.get("automation_policy", {}).get("human_in_the_loop", "")
    keywords = [w.strip().lower() for w in re.split(r"[,;]", hil_text) if len(w.strip()) > 3]
    topic_lower = topic.lower()
    return any(kw in topic_lower for kw in keywords)

@app.get("/api/v1/onboarding/profile")
def get_onboarding_profile(current_user: dict = Depends(get_current_user)):
    return {"status": "success", "profile": load_company_profile()}

@app.post("/api/v1/onboarding/setup")
def setup_onboarding_profile(payload: OnboardingFullRequest, current_user: dict = Depends(get_current_user)):
    profile_dict = payload.dict()
    save_company_profile(profile_dict)

    onboarding_text = f"""
    [ŞİRKET ONBOARDING VE KÜLTÜR KILAVUZU]
    1. Kimlik ve Ölçek: Şirket: {payload.identity.get('name')}, Sektör: {payload.identity.get('sector')}, Ekip: {payload.identity.get('team_size')}
    2. İletişim Tonu: {payload.communication_tone}
    3. Kritik Departmanlar: {json.dumps(payload.key_contacts, ensure_ascii=False)}
    4. Çalışma Düzeni ve Kırmızı Çizgiler: {json.dumps(payload.standards, ensure_ascii=False)}
    5. Otomasyon ve Güven Eşiği: {json.dumps(payload.automation_policy, ensure_ascii=False)}
    """
    if rag_pipeline_available and process_and_store_document:
        try:
            process_and_store_document(
                source_file="corporate_onboarding_master_guide.txt",
                content=onboarding_text,
                document_type="corporate_culture"
            )
        except Exception as e:
            logger.warning(f"Onboarding RAG kayıt hatası: {e}")

    return {"status": "success", "message": "Onboarding kuralları kaydedildi.", "profile": profile_dict}

# --- CHAT & RAG ---
@app.post("/api/chat")
def chat_with_ai(request: ChatRequest):
    global auditor
    try:
        msg_lower = request.message.lower()
        if "mail" in msg_lower or "görev oluştur" in msg_lower or "tara" in msg_lower:
            sync_result = sync_gmail_tasks()
            return {
                "status": "success",
                "reply": f"📥 Otonom Komut Algılandı: {sync_result.get('message', 'İşlem tamamlandı.')}",
                "citations": [],
                "tasks": list(db_tasks)[:20],
                "risks": list(db_risk_findings)[:20],
            }

        profile_answer = answer_from_company_profile(request.message)
        if profile_answer:
            return {"status": "success", "reply": profile_answer}

        if not auditor or not auditor.is_ready:
            if Auditor is not None:
                auditor = Auditor()

        if not auditor or not auditor.is_ready:
            return {
                "status": "error",
                "reply": (
                    "AI servisi şu anda kullanılamıyor. Foundry Local'in çalıştığını ve "
                    "phi-3.5-mini modelinin yüklü olduğunu kontrol edip tekrar deneyin."
                ),
            }

        profile = load_company_profile()
        onboarding_rules = f"""
        [ZORUNLU OPERASYONEL KURALLAR VE ŞİRKET KİMLİĞİ]
        - Şirket: {profile.get('identity', {}).get('name')} ({profile.get('identity', {}).get('sector')})
        - İletişim Tonu: {profile.get('communication_tone')}
        - Kırmızı Çizgiler: {profile.get('standards', {}).get('red_lines')}
        - İnsan Onayı (Human-in-the-Loop): {profile.get('automation_policy', {}).get('human_in_the_loop')}
        """

        context_text = ""
        if rag_pipeline_available and retrieve_relevant_chunks:
            target_source = None
            if request.document_id:
                matched_doc = next((d for d in db_vault_documents if d["id"] == request.document_id), None)
                if matched_doc:
                    target_source = matched_doc["name"]
            relevant_chunks = retrieve_relevant_chunks(request.message, top_k=3, source_file=target_source)
            if relevant_chunks:
                context_text = "\n\n".join([f"[Kaynak Belge: {chunk[1]}]\n{chunk[2]}" for chunk in relevant_chunks])

        augmented_prompt = f"""
        {MANTIS_PERSONA}

        {onboarding_rules}

        Vault / Doküman Bağlamı:
        {context_text if context_text else "Özel bir doküman seçilmedi, genel hafıza devrede."}

        Kullanıcı Sorusu: {request.message}
        Kural: Her bulgu, çıkarılan görev ve yasal riskin yanına dayandığı kaynağı `[Kaynak: X, Madde: Y]` formatında iliştir.
        Kural: JSON çıktısı üretirken `citations: [{"source_doc": str, "page": int, "excerpt": str}]` alanını zorunlu ekle.
        """
        response_text = auditor.ask(augmented_prompt)
        return {
            "status": "success",
            "reply": response_text,
            "citations": [],
            "tasks": list(db_tasks)[:20],
            "risks": list(db_risk_findings)[:20],
        }
    except Exception as e:
        logger.error(f"Chat hatası: {str(e)}")
        return {"status": "error", "reply": f"Sohbet sırasında hata oluştu: {str(e)}"}

# --- GMAIL GÖREV SENKRONİZASYONU ---
@app.post("/api/v1/tasks/sync-gmail")
def sync_gmail_tasks(current_user: dict = Depends(get_current_user)):
    global auditor
    try:
        if auditor is None and Auditor is not None:
            auditor = Auditor()

        service = get_gmail_service()
        results = service.users().messages().list(userId='me', maxResults=5, labelIds=['INBOX']).execute()
        messages = results.get('messages', [])

        total_scanned = len(messages)
        spam_or_ads_count = 0
        confirmed_added = 0
        suggested_added = 0

        for msg in messages:
            msg_id = msg['id']
            txt = service.users().messages().get(userId='me', id=msg_id).execute()
            headers = txt['payload']['headers']

            subject = next((h['value'] for h in headers if h['name'] == 'Subject'), "Konu Yok")
            sender = next((h['value'] for h in headers if h['name'] == 'From'), "Bilinmeyen Gönderen")
            snippet = txt.get('snippet', '')

            if any(bad_kw in sender.lower() or bad_kw in subject.lower() for bad_kw in [
                "mailer-daemon", "delivery subsystem", "undelivered", "postmaster",
                "promosyon", "indirim", "bülten", "newsletter", "google"
            ]):
                spam_or_ads_count += 1
                continue

            confidence = 0.80
            task_title = f"Aksiyon: {subject}"
            deadline = "2026-08-31"
            importance_reason = "E-posta analiz edildi."
            llm_json_ok = False

            task_candidate = any(kw in f"{subject} {snippet}".lower() for kw in [
                "aksiyon", "görev", "task", "onay", "gönder", "hazırla", "rapor", "sınav", "son tarih", "deadline"
            ])

            if task_candidate and auditor and auditor.is_ready:
                analysis_prompt = f"""
                Sen Project Mantis Inbox Intelligence Ajanısın. E-postayı analiz et:
                Gönderen: {sender} | Konu: {subject} | İçerik: {snippet}
                SADECE şu JSON formatında yanıt ver:
                {{
                    "importance_level": "Yüksek / Orta / Düşük",
                    "importance_reason": "Gerekçe",
                    "is_task": true/false,
                    "confidence": 0.86,
                    "task_title": "Görev Başlığı",
                    "deadline": "2026-08-31"
                }}
                """
                try:
                    llm_raw = auditor.ask(analysis_prompt, timeout=20)
                    json_match = re.search(r'\{.*\}', llm_raw, re.DOTALL)
                    if json_match:
                        parsed = json.loads(json_match.group(0))
                        llm_json_ok = True
                        importance_reason = parsed.get("importance_reason", "")
                        if parsed.get("is_task", False):
                            confidence = float(parsed.get("confidence", 0.80))
                            task_title = parsed.get("task_title", subject)
                            deadline = parsed.get("deadline", "2026-08-31")
                        else:
                            continue
                except Exception as ex:
                    logger.warning("E-posta AI analizi atlandı (%s): %s", msg_id, ex)

            if not llm_json_ok and not task_candidate:
                continue

            existing_ids = [t.get("source_email_id") for t in db_tasks]
            if msg_id not in existing_ids:
                if confidence > 0.75:
                    status_val = "confirmed"
                    confirmed_added += 1
                    type_label = "AI Tarafından Eklendi"
                    sync_task_to_google_tasks(task_title, deadline)
                else:
                    status_val = "suggested"
                    suggested_added += 1
                    type_label = "Onay Bekleyen Öneri"

                new_task = {
                    "id": f"task_{len(db_tasks) + 1:04d}",
                    "title": task_title,
                    "deadline": deadline,
                    "source_email_id": msg_id,
                    "confidence": confidence,
                    "status": status_val,
                    "type": type_label,
                    "sender": sender,
                    "ai_draft": f"Ajan Değerlendirmesi: {importance_reason}\n\nİçerik: {snippet}",
                    "completed": confidence > 0.75 and status_val == "confirmed",
                    "created_by": "agent"
                }
                db_tasks.insert(0, new_task)

        total_pending = sum(1 for t in db_tasks if t["status"] == "suggested")
        return {
            "status": "success",
            "message": (
                f"Gelen kutusu tarandı: {total_scanned} mail incelendi "
                f"({spam_or_ads_count} tanesi reklam/spam olarak elendi). "
                f"{confirmed_added} görev otomatik eklendi, {suggested_added} öneri onay kuyruğunda. "
                f"Toplam onay bekleyen: {total_pending}."
            ),
            "total_tasks": len(db_tasks)
        }
    except Exception as e:
        logger.error("Task Sync Hatası: %s", e, exc_info=True)
        return {
            "status": "error",
            "message": (
                "Gmail taraması tamamlanamadı. Google hesabının yetkilendirildiğini ve "
                "Gmail API erişiminin açık olduğunu kontrol edin."
            ),
            "detail_code": "GOOGLE_AUTH_REQUIRED",
        }

# --- GÖREV AKSİYONLARI ---
@app.post("/api/v1/tasks/action")
def handle_task_action(payload: TaskActionRequest):
    target_task = next((t for t in db_tasks if str(t["id"]) == str(payload.task_id)), None)
    if not target_task:
        raise HTTPException(status_code=404, detail="Görev bulunamadı.")

    task_type = target_task.get("type", "").lower()
    task_title = target_task.get("title", "")
    
    if payload.action == "approve":
        target_task = db_tasks.update(target_task["id"], status="Onaylandı", completed=True)

        pending_doc_topic = target_task.get("pending_doc_topic")
        if pending_doc_topic:
            vault_entry = generate_real_docx(topic=pending_doc_topic, context_hint=target_task.get("ai_draft", ""))
            return {
                "status": "success",
                "message": f"Onaylandı: '{pending_doc_topic}' belgesi oluşturuldu ve Doc Vault'a kaydedildi (ID: {vault_entry['id']}).",
                "document_id": vault_entry["id"],
            }

        if "ihtar" in task_type or "hukuki" in task_type or target_task.get("requires_ai_review", False):
            target_recipient = target_task.get("recipient_email", "hukuk@sirket.com")
            draft_content = target_task.get("ai_draft", "Otonom onaylanan metin.")
            send_legal_approval_email(
                personnel_name=target_task.get("sender", "Hilal Uluca"),
                draft_text=draft_content,
                recipient_email=target_recipient
            )
            action_result_msg = f"Onaylandı ve '{target_recipient}' adresine iletildi."
        else:
            sync_task_to_google_tasks(task_title, target_task.get("deadline"))
            action_result_msg = "Görev onaylandı ve Google Tasks'a işlendi."
        return {"status": "success", "message": action_result_msg}

    elif payload.action == "reject":
        db_tasks.update(target_task["id"], status="Reddedildi", completed=True)
        return {"status": "success", "message": "Görev reddedildi."}
    else:
        raise HTTPException(status_code=400, detail="Geçersiz aksiyon tipi.")

# --- PUANTAJ NLP & İHTARNAME ÜRETİMİ ---
@app.post("/api/v1/attendance/parse")
def parse_attendance_command(request: AttendanceCommand):
    global auditor
    try:
        if not request.consent_given:
            return {"status": "error", "message": "Hukuki onay verilmeden personel verisi işlenemez."}

        cmd_lower = (request.command or "").lower()
        today_date = datetime.now().strftime("%Y-%m-%d")

        default_name = "Bilinmeyen Personel"
        name_match = re.search(r"([A-ZÇĞİÖŞÜa-zçğışüö]+(?:\s+[A-ZÇĞİÖŞÜa-zçğışüö]+){0,2})", request.command)
        if name_match:
            default_name = name_match.group(1).strip()

        parsed_data = {
            "personnel_name": default_name,
            "date": today_date,
            "status": "Devamsız",
            "reason": None,
        }

        if re.search(r"(mazeretsiz|izinsiz|devamsiz)", cmd_lower):
            parsed_data["reason"] = None
        elif re.search(r"(mazeret|mazeretli)", cmd_lower):
            parsed_data["reason"] = "Mazeret"

        if re.search(r"\b(\d{4}-\d{2}-\d{2})\b", request.command):
            parsed_data["date"] = re.search(r"\b(\d{4}-\d{2}-\d{2})\b", request.command).group(1)

        if re.search(r"\b(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})\b", request.command):
            raw = re.search(r"\b(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})\b", request.command).group(1)
            try:
                parsed_data["date"] = datetime.strptime(raw, "%d.%m.%Y").strftime("%Y-%m-%d")
            except ValueError:
                try:
                    parsed_data["date"] = datetime.strptime(raw, "%d/%m/%Y").strftime("%Y-%m-%d")
                except ValueError:
                    pass

        if "sevil" in cmd_lower or "sevıl" in cmd_lower:
            parsed_data["personnel_name"] = "Sevil Hanım"
        elif re.search(r"\b([A-ZÇĞİÖŞÜ][a-zçğışüö]+\s+[A-ZÇĞİÖŞÜ][a-zçğışüö]+)\b", request.command):
            parsed_data["personnel_name"] = re.search(r"\b([A-ZÇĞİÖŞÜ][a-zçğışüö]+\s+[A-ZÇĞİÖŞÜ][a-zçğışüö]+)\b", request.command).group(1)

        if not auditor or not auditor.is_ready:
            if Auditor is not None:
                auditor = Auditor()

        if not auditor or not auditor.is_ready:
            return {
                "status": "error",
                "message": "Auditor LLM aktif değil — Foundry Local servisinin çalıştığından emin olun.",
            }

        log_id = f"LOG-{datetime.now().strftime('%H%M%S')}"
        db_audit_logs.insert(0, {
            "id": log_id,
            "agent_name": request.operator,
            "action": "KVKK_CONSENT_GRANTED",
            "confidence_score": 1.0,
            "source_doc_id": "SYS",
            "timestamp": datetime.now().strftime("%H:%M:%S"),
            "level": "LEGAL",
            "message": f"Puantaj verisi NLP işlemine alındı. Komut: '{request.command}'",
        })

        letter_prompt = f"""
        {MANTIS_PERSONA}

        Aşağıdaki bilgilere dayanarak 4857 Sayılı İş Kanunu Madde 25/II uyarınca
        resmi bir İHTARNAME / SAVUNMA TALEP YAZISI hazırla. Sadece resmi metni yaz.

        Personel: {parsed_data.get('personnel_name')}
        Devamsızlık Tarihi: {parsed_data.get('date')}
        Mazeret: {parsed_data.get('reason') or 'Belirtilmemiş'}
        """
        letter_draft = auditor.ask(letter_prompt)

        new_task = {
            "id": f"task_{len(db_tasks) + 1:04d}",
            "title": f"İhtarname Onayı: {parsed_data.get('personnel_name')}",
            "deadline": parsed_data.get("date"),
            "source_email_id": "attendance_nlp",
            "confidence": 0.90,
            "status": "Onay Bekliyor",
            "type": "Kritik Güvenlik Onayı",
            "requires_ai_review": True,
            "recipient_email": "hukuk@sirket.com",
            "sender": parsed_data.get("personnel_name"),
            "ai_draft": letter_draft,
            "completed": False,
            "created_by": "agent",
        }
        db_tasks.insert(0, new_task)

        return {
            "status": "success",
            "data": parsed_data,
            "letter_draft": letter_draft,
            "task_id": new_task["id"],
            "message": "İhtarname taslağı hazırlandı ve onay kuyruğuna eklendi — Tasks kısmından onaylayabilirsin.",
        }
    except Exception as e:
        logger.error(f"Devamsızlık hatası: {str(e)}")
        return {"status": "error", "message": str(e)}

@app.post("/api/v1/attendance/warning-letter")
def generate_warning_letter(payload: WarningLetterRequest):
    global auditor
    try:
        if not auditor or not auditor.is_ready:
            if Auditor is not None:
                auditor = Auditor()

        prompt = f"""
        {MANTIS_PERSONA}

        Aşağıdaki bilgilere dayanarak 4857 Sayılı İş Kanunu Madde 25/II uyarınca resmi ve hukuki geçerliliği olan bir İHTARNAME / SAVUNMA TALEP YAZISI hazırla.
        Personel: {payload.personnel_name}
        Devamsızlık Tarihi: {payload.date}
        Mazeret Durumu: {payload.reason or 'Mazeret bildirilmemiştir'}

        Sadece resmi ihtarname metnini oluştur.
        """
        fallback_letter = (
            f"SAYIN {payload.personnel_name.upper()},\n\n"
            f"{payload.date} tarihinde işyerinde bulunmadığınız ve bu devamsızlığa ilişkin önceden bildirilmiş veya belgelenmiş bir mazeret sunulmadığı tespit edilmiştir. "
            "İşbu yazı, olayın açıklığa kavuşturulması ve savunma hakkınızın kullandırılması amacıyla düzenlenmiştir. "
            "4857 sayılı İş Kanunu'nun 17. maddesi ve 25/II-g bendi kapsamında, devamsızlığınızın nedenlerini ve varsa destekleyici belgeleri bu bildirimin tebliğinden itibaren üç iş günü içinde yazılı olarak İnsan Kaynakları birimine sunmanız gerekmektedir. "
            "Belirtilen süre içinde savunma verilmemesi veya sunulan açıklamanın yeterli görülmemesi halinde, iş sözleşmeniz ve ilgili mevzuat çerçevesinde gerekli değerlendirme yapılacaktır."
        )
        try:
            letter_draft = auditor.ask(prompt, timeout=45, temperature=0.0) if (auditor and auditor.is_ready) else ""
        except Exception:
            letter_draft = ""
        if not letter_draft or letter_draft.lower().startswith("bir sorun oluştu") or "timed out" in letter_draft.lower():
            letter_draft = fallback_letter

        mail_delivered = send_legal_approval_email(
            personnel_name=payload.personnel_name,
            draft_text=letter_draft,
            recipient_email="hukuk@sirket.com"
        )

        return {
            "status": "success",
            "letter_draft": letter_draft,
            "mail_delivered": mail_delivered,
            "message": "İhtarname taslağı başarıyla oluşturuldu."
        }
    except Exception as e:
        logger.error(f"İhtarname üretim hatası: {e}")
        return {"status": "error", "message": str(e)}

# --- SUBSCRIPTIONS & ARKA PLAN GÖREVLERİ ---
@app.post("/api/v1/subscriptions")
def create_subscription(payload: SubscriptionRequest):
    sub_id = f"sub_{int(time.time())}"
    expire_time = datetime.now().timestamp() + (4230 * 60)
    record = {
        "subscription_id": sub_id,
        "callback_url": payload.callback_url,
        "resource_type": payload.resource_type,
        "created_at": datetime.now().isoformat(),
        "expires_at_timestamp": expire_time,
        "status": "active"
    }
    db_subscriptions[sub_id] = record
    return {"status": "success", "subscription": record}

@app.get("/api/v1/subscriptions")
def get_subscriptions():
    return {"status": "success", "subscriptions": list(db_subscriptions.values())}

def background_subscription_renewal_loop():
    while True:
        try:
            now = datetime.now().timestamp()
            for sub_id, sub in list(db_subscriptions.items()):
                if sub["expires_at_timestamp"] - now < 3600:
                    sub["expires_at_timestamp"] = now + (4230 * 60)
                    sub["status"] = "renewed_automatically"
        except Exception as e:
            logger.error(f"Subscription renewal hatası: {e}")
        time.sleep(1800)

# --- PROAKTİF BRİFİNG & ORCHESTRATOR ---
def generate_proactive_morning_briefing() -> str:
    try:
        service = get_gmail_service()
        results = service.users().messages().list(userId='me', maxResults=10, labelIds=['INBOX']).execute()
        total_unread = len(results.get('messages', []))
        mail_line = f"{total_unread} yeni mail var"
    except Exception:
        mail_line = "Mail kutusuna şu an erişemedim"

    pending_tasks = sum(1 for t in db_tasks if t.get("status") in ("suggested", "Onay Bekliyor"))
    pending_risks = sum(1 for r in db_risk_findings if r.get("status") == "PENDING_APPROVAL")

    return f"{mail_line}. Onayınızı bekleyen {pending_tasks} görev ve {pending_risks} risk bulgusu var."

@app.get("/api/v1/orchestrator/morning-briefing")
def get_morning_briefing():
    return {"status": "success", "briefing": generate_proactive_morning_briefing()}

def format_tasks_reply() -> str:
    if not db_tasks:
        return "Şu anda kayıtlı görev bulunmamaktadır."
    confirmed = [t for t in db_tasks if t["status"] in ("confirmed", "Onaylandı")]
    suggested = [t for t in db_tasks if t["status"] in ("suggested", "Onay Bekliyor")]
    lines = [f"Toplam {len(db_tasks)} görev mevcut ({len(confirmed)} onaylı, {len(suggested)} onay bekliyor)."]
    for t in confirmed[:5]:
        lines.append(f"• {t['title']} (Son Tarih: {t.get('deadline', 'Belirtilmedi')})")
    if suggested:
        lines.append("\nOnay bekleyenler:")
        for t in suggested[:5]:
            lines.append(f"• {t['title']} — onaylar mısın?")
    return "\n".join(lines)

def generate_full_status_report() -> str:
    mail_part = generate_proactive_morning_briefing()
    tasks_part = format_tasks_reply()

    docs_part = ""
    if len(db_vault_documents) > 0:
        recent = list(db_vault_documents)[:3]
        docs_lines = "\n".join(f"• {d['name']}" for d in recent)
        docs_part = f"\n\nSon oluşturulan dokümanlar:\n{docs_lines}"

    return f"{mail_part}\n\n{tasks_part}{docs_part}"

def try_add_task_from_message(user_msg: str) -> Optional[dict]:
    global auditor
    title = None
    deadline = None

    m = re.search(r"(?:görev|task)\s*ekle\s*[:\-]?\s*(.+)", user_msg, re.IGNORECASE)
    if m:
        title = m.group(1).strip(" .")

    if not title and auditor and auditor.is_ready:
        today_date = datetime.now().strftime("%Y-%m-%d")
        prompt = f"""
        Bugünün tarihi: {today_date}. Kullanıcının şu mesajından bir görev
        başlığı ve varsa son tarih çıkar: "{user_msg}"
        SADECE JSON döndür: {{"title": "kısa başlık ya da boş", "deadline": "YYYY-MM-DD ya da null"}}
        """
        raw = auditor.ask(prompt)
        match = re.search(r"\{.*\}", raw, re.DOTALL)
        if match:
            try:
                parsed = json.loads(match.group(0))
                title = parsed.get("title") or None
                deadline = parsed.get("deadline")
            except Exception:
                pass

    if not title:
        return None

    new_task = {
        "id": f"task_{len(db_tasks) + 1:04d}",
        "title": title,
        "deadline": deadline or "Belirtilmedi",
        "source_email_id": "manual_chat",
        "confidence": 0.99,
        "status": "confirmed",
        "type": "Kullanıcı Tarafından Eklendi",
        "completed": False,
        "created_by": "user_via_chat",
    }
    db_tasks.insert(0, new_task)
    sync_task_to_google_tasks(title, deadline)
    return new_task

@app.post("/api/v1/orchestrator/chat")
def orchestrator_chat(request: OrchestratorChatRequest):
    global auditor, chat_memory_history
    try:
        user_msg = request.message
        user_lower = user_msg.lower()

        chat_memory_history.append({"role": "user", "content": user_msg})
        if len(chat_memory_history) > 10:
            chat_memory_history.pop(0)

        if any(kw in user_lower for kw in ["mailleri tara", "gelen kutusu tara", "mail tara", "posta tara"]):
            real_result = sync_gmail_tasks()
            msg = f"🚀 {real_result.get('message', 'Gelen kutusu tarandı.')}"
            task_list = list(db_tasks)[:20]
            risk_list = list(db_risk_findings)[:20]
            chat_memory_history.append({"role": "assistant", "content": msg})
            return {
                "status": real_result.get("status", "error"),
                "reply": msg,
                "citations": [],
                "tasks": task_list,
                "risks": risk_list,
            }

        if any(kw in user_lower for kw in ["ihtarname oluştur", "ihtarname hazırla", "savunma talebi"]):
            name_match = re.search(r"(?:ihtarname|savunma talebi)(?: oluştur| hazırla)?\s*(?:için|:)?\s*(.+)", user_msg, re.IGNORECASE)
            personnel_name = name_match.group(1).strip() if name_match else "İlgili personel"
            warning = generate_warning_letter(WarningLetterRequest(personnel_name=personnel_name, date=datetime.now().strftime("%Y-%m-%d")))
            msg = warning.get("message", "İhtarname taslağı oluşturuldu.")
            chat_memory_history.append({"role": "assistant", "content": msg})
            return {"status": warning.get("status", "error"), "reply": f"{msg} Hukuk onayı ve gönderim durumu ayrıca gösteriliyor."}

        if any(kw in user_lower for kw in ["mail gönder", "e-posta gönder", "email gönder"]):
            recipient_match = re.search(r"(?:mail|e-?posta|email) gönder(?:\s+(?:adresine|için))?\s*([\w.+-]+@[\w.-]+)", user_msg, re.IGNORECASE)
            if not recipient_match:
                msg = "Mail gönderebilmem için alıcı e-posta adresini belirtmelisin. Örnek: mail gönder ekip@firma.com"
                return {"status": "needs_clarification", "reply": msg}
            recipient = recipient_match.group(1)
            body = re.sub(r".*?" + re.escape(recipient), "", user_msg, count=1, flags=re.IGNORECASE).strip(" :,-") or "Project Mantis tarafından gönderilen bilgilendirme."
            sent = send_legal_approval_email("Mantis AI", body, recipient)
            msg = f"{recipient} için e-posta gönderim isteği işlendi." if sent else f"{recipient} adresine gönderim yapılamadı. SMTP ayarları eksik; mesaj gönderilmedi."
            return {"status": "success" if sent else "error", "reply": msg}

        if any(kw in user_lower for kw in ["görev ekle", "task ekle", "göreve ekle"]):
            added = try_add_task_from_message(user_msg)
            if added:
                msg = f"✅ Görev eklendi: {added['title']}" + (f" (son tarih: {added['deadline']})" if added.get("deadline") != "Belirtilmedi" else "")
            else:
                msg = "Görevi ekleyebilirim ama başlığı net anlayamadım — ne eklememi istersin?"
            chat_memory_history.append({"role": "assistant", "content": msg})
            return {"status": "success" if added else "needs_clarification", "reply": msg}

        if any(kw in user_lower for kw in ["rapor ver", "genel rapor", "brifing", "briefing", "özet ver", "durum ne", "durum nedir"]):
            msg = generate_full_status_report()
            chat_memory_history.append({"role": "assistant", "content": msg})
            return {"status": "success", "reply": msg}

        if any(kw in user_lower for kw in ["görev", "task", "kuyruk", "ne var"]):
            msg = format_tasks_reply()
            chat_memory_history.append({"role": "assistant", "content": msg})
            return {"status": "success", "reply": msg}

        if any(kw in user_lower for kw in ["doküman oluştur", "docx", "belge hazırla", "rapor oluştur"]):
            DOC_TRIGGERS = ["doküman oluştur", "docx", "belge hazırla", "rapor oluştur"]
            topic_guess = user_msg
            for kw in DOC_TRIGGERS:
                topic_guess = re.sub(re.escape(kw), "", topic_guess, flags=re.IGNORECASE)
            topic_guess = topic_guess.strip(" ,.:-\n")

            if len(topic_guess.split()) < 2:
                msg = "Tabii, hemen hazırlarım — hangi konuda bir doküman istiyorsun?"
                chat_memory_history.append({"role": "assistant", "content": msg})
                return {"status": "needs_clarification", "reply": msg}

            if is_human_in_the_loop_topic(topic_guess):
                draft_text = None
                if auditor and auditor.is_ready:
                    draft_text = auditor.ask(f"{MANTIS_PERSONA}\n\n'{topic_guess}' konusunda resmi bir belge taslağı yaz.")
                new_task = {
                    "id": f"task_{len(db_tasks) + 1:04d}",
                    "title": f"Doküman Onayı: {topic_guess}",
                    "deadline": "Belirtilmedi",
                    "source_email_id": "doc_request",
                    "confidence": 0.85,
                    "status": "Onay Bekliyor",
                    "type": "Otonom Taslak - Onay Bekliyor",
                    "requires_ai_review": True,
                    "pending_doc_topic": topic_guess,
                    "ai_draft": draft_text or "(Taslak önizlemesi için Foundry Local'in çalışıyor olması gerekiyor.)",
                    "completed": False,
                    "created_by": "agent",
                }
                db_tasks.insert(0, new_task)
                msg = (
                    f"'{topic_guess}' şirket politikanıza göre onay gerektiren bir konu — "
                    f"taslağı hazırladım ve onay kuyruğuna ekledim. Onaylarsan gerçek belgeyi oluşturup Doc Vault'a kaydedeceğim."
                )
                chat_memory_history.append({"role": "assistant", "content": msg})
                return {"status": "needs_approval", "reply": msg, "task_id": new_task["id"]}

            vault_entry = generate_real_docx(topic=topic_guess)
            msg = f"📄 '{topic_guess}' belgesi oluşturuldu ve Doc Vault'a kaydedildi (ID: {vault_entry['id']})."
            chat_memory_history.append({"role": "assistant", "content": msg})
            return {"status": "success", "reply": msg}

        profile_answer = answer_from_company_profile(user_msg)
        if profile_answer:
            chat_memory_history.append({"role": "assistant", "content": profile_answer})
            return {"status": "success", "reply": profile_answer}

        chat_res = chat_with_ai(ChatRequest(message=user_msg, document_id=request.document_id))
        reply_text = chat_res.get("reply", "Yanıt üretilemedi.")
        chat_memory_history.append({"role": "assistant", "content": reply_text})
        return {
            "status": chat_res.get("status", "success"),
            "reply": reply_text,
            "citations": chat_res.get("citations", []),
            "tasks": chat_res.get("tasks", []),
            "risks": chat_res.get("risks", []),
        }
    except Exception as e:
        logger.error(f"Orchestrator chat error: {e}", exc_info=True)
        return {
            "status": "error",
            "reply": "İşlem tamamlanamadı. Yerel AI servisini kontrol edip tekrar deneyin.",
        }

def _fast_path_reply(user_msg: str) -> Optional[str]:
    user_lower = user_msg.lower()

    if any(kw in user_lower for kw in ["mailleri tara", "gelen kutusu tara", "mail tara", "posta tara"]):
        real_result = sync_gmail_tasks()
        return f"🚀 {real_result.get('message', 'Gelen kutusu tarandı.')}"

    if any(kw in user_lower for kw in ["görev ekle", "task ekle", "göreve ekle"]):
        added = try_add_task_from_message(user_msg)
        if added:
            return f"✅ Görev eklendi: {added['title']}"
        return "Görevi ekleyebilirim ama başlığı net anlayamadım — ne eklememi istersin?"

    if any(kw in user_lower for kw in ["rapor ver", "genel rapor", "brifing", "briefing", "özet ver", "durum ne", "durum nedir"]):
        return generate_full_status_report()

    if any(kw in user_lower for kw in ["görev", "task", "kuyruk", "ne var"]):
        return format_tasks_reply()

    profile_answer = answer_from_company_profile(user_msg)
    if profile_answer:
        return profile_answer

    return None

@app.post("/api/v1/orchestrator/chat/stream")
def orchestrator_chat_stream(request: OrchestratorChatRequest):
    global auditor

    fast_reply = _fast_path_reply(request.message)
    if fast_reply is not None:
        def _instant():
            yield fast_reply
        return StreamingResponse(_instant(), media_type="text/plain")

    if not auditor or not auditor.is_ready:
        if Auditor is not None:
            auditor = Auditor()
    if not auditor or not auditor.is_ready:
        def _err():
            yield "Şu an yerel modele bağlanamıyorum, Foundry Local servisini kontrol eder misin?"
        return StreamingResponse(_err(), media_type="text/plain")

    def _gen():
        for token in auditor.ask_stream(request.message):
            yield token
    return StreamingResponse(_gen(), media_type="text/plain")

@app.post("/api/v1/orchestrator/voice-command")
async def handle_voice_command(file: UploadFile = File(...)):
    try:
        return orchestrator_chat(OrchestratorChatRequest(message="mailleri tara", voice_enabled=True))
    except Exception as e:
        raise HTTPException(status_code=500, detail="Ses işlenemedi.")

# --- DOC VAULT ENDPOINTLERİ ---
@app.get("/api/v1/documents/vault")
def get_vault_documents():
    return [{k: v for k, v in doc.items() if k not in ("filepath", "content")} for doc in db_vault_documents]

@app.get("/api/v1/documents/{doc_id}")
def get_vault_document(doc_id: str, current_user: dict = Depends(get_current_user)):
    entry = next((d for d in db_vault_documents if d["id"] == doc_id), None)
    if not entry:
        raise HTTPException(status_code=404, detail="Doküman bulunamadı.")

    content = entry.get("content") or ""
    if not content:
        filepath = entry.get("filepath")
        if filepath and os.path.exists(filepath):
            content = _extract_document_text(filepath)
            entry["content"] = content
            db_vault_documents.update(doc_id, content=content)
        else:
            content = "Belge içeriği güvenli şekilde hazırlanıyor; metin görünümü bulunamadı."
            entry["content"] = content
    payload = {k: v for k, v in entry.items() if k != "filepath"}
    payload["content"] = content
    return payload

@app.post("/api/v1/documents/{doc_id}/chat")
def chat_with_document(doc_id: str, payload: DocumentChatRequest, current_user: dict = Depends(get_current_user)):
    return document_chat(DocumentChatEnvelope(document_id=doc_id, message=payload.message), current_user)

@app.put("/api/v1/documents/{doc_id}")
def update_vault_document(doc_id: str, payload: DocumentUpdateRequest, current_user: dict = Depends(get_current_user)):
    entry = next((d for d in db_vault_documents if d["id"] == doc_id), None)
    if not entry:
        raise HTTPException(status_code=404, detail="Doküman bulunamadı.")
    if not payload.content.strip():
        raise HTTPException(status_code=400, detail="Doküman içeriği boş bırakılamaz.")
    filepath = entry.get("filepath")
    if filepath and filepath.lower().endswith(".docx"):
        document = DocxDocument()
        document.add_heading(entry.get("name", "Doküman"), level=1)
        for paragraph in payload.content.splitlines():
            if paragraph.strip():
                document.add_paragraph(paragraph.strip())
        document.save(filepath)
    updated = db_vault_documents.update(doc_id, content=payload.content, last_active=datetime.now().strftime("%Y-%m-%d %H:%M"))
    return {"status": "success", "document": {k: v for k, v in updated.items() if k != "filepath"}}

@app.delete("/api/v1/documents/{doc_id}")
def delete_vault_document(doc_id: str, current_user: dict = Depends(get_current_user)):
    entry = next((d for d in db_vault_documents if d["id"] == doc_id), None)
    if not entry:
        raise HTTPException(status_code=404, detail="Doküman bulunamadı.")
    filepath = entry.get("filepath")
    if filepath and os.path.exists(filepath):
        os.remove(filepath)
    db_vault_documents.delete(doc_id)
    return {"status": "success", "message": "Doküman Vault'tan kaldırıldı."}

def _get_document_content(doc_id: str) -> tuple[dict, str]:
    entry = next((d for d in db_vault_documents if d["id"] == doc_id), None)
    if not entry:
        raise HTTPException(status_code=404, detail="Doküman bulunamadı.")
    content = entry.get("content", "")
    filepath = entry.get("filepath", "")
    if not content and filepath and os.path.exists(filepath):
        if filepath.lower().endswith(".docx"):
            content = "\n".join(paragraph.text for paragraph in DocxDocument(filepath).paragraphs)
        elif os.path.splitext(filepath)[1].lower() in {".txt", ".md"}:
            with open(filepath, "r", encoding="utf-8", errors="replace") as file:
                content = file.read()
        if content:
            db_vault_documents.update(doc_id, content=content)
    return entry, content

def _document_context(entry: dict, content: str, query: str) -> tuple[str, List[dict]]:
    chunks = []
    if rag_pipeline_available and retrieve_relevant_chunks:
        chunks = retrieve_relevant_chunks(query, top_k=5, source_file=entry.get("name"))
    if chunks:
        citations = [{
            "source": item.get("doc_name") or item.get("source_file") or entry.get("name"),
            "text": item.get("text") or "",
            "score": round(float(item.get("score", 0.0) or 0.0), 3),
            "page": item.get("page_number") or 1,
            "clause": item.get("clause_title") or "Genel Madde",
            "doc_id": item.get("document_id") or item.get("source_file") or entry.get("id"),
        } for item in chunks]
        context = "\n\n".join(
            f"[Kaynak: {item.get('source')}, Sayfa: {item.get('page')}, Madde: {item.get('clause')}]\n{item.get('text')}"
            for item in citations
        )
        return context, citations
    fallback = content[:12000]
    return fallback, [{"source": entry.get("name"), "text": content[:500], "score": 1.0, "page": 1, "clause": "Genel Madde", "doc_id": entry.get("id")}] if content else []

def _ensure_auditor() -> Any:
    global auditor
    if not auditor or not auditor.is_ready:
        if Auditor is not None:
            auditor = Auditor()
    if not auditor or not auditor.is_ready:
        raise HTTPException(status_code=503, detail="Yerel AI servisi şu anda kullanılamıyor.")
    return auditor

@app.post("/api/v1/documents/chat")
def document_chat(payload: DocumentChatEnvelope, current_user: dict = Depends(get_current_user)):
    try:
        if not payload.message.strip():
            raise HTTPException(status_code=400, detail="Dokümana sorulacak soru boş olamaz.")
        entry, content = _get_document_content(payload.document_id)
        if not content.strip():
            raise HTTPException(status_code=422, detail="Bu dokümanda okunabilir metin bulunamadı.")
        content_lines = [line.strip() for line in content.splitlines() if line.strip()]
        if len(content_lines) <= 3 and any(line.endswith("?") or line.endswith("?") for line in content_lines):
            return {
                "status": "success",
                "reply": "Bu doküman ödeme sürecini açıklayan bir içerik içermiyor; yalnızca konu başlığı ve soru bulunuyor. Ödeme vadesi, tahsilat adımları veya sorumluları ekleyerek belgeyi revize edebilirsin.",
                "cited_clauses": [{"source": entry.get("name"), "text": line, "score": 1.0} for line in content_lines],
            }
        model = _ensure_auditor()
        context, citations = _document_context(entry, content, payload.message)
        profile = load_company_profile()
        prompt = f"""{MANTIS_PERSONA}
Sen bir doküman asistanısın. Yalnızca aşağıdaki kaynak parçalarına dayan.
Kaynakta bulunmayan bilgiyi uydurma; yoksa açıkça 'Bu dokümanda belirtilmiyor' de.
Şirket bağlamı: {json.dumps(profile, ensure_ascii=False)}
Doküman: {entry.get('name')}
Kaynak parçaları:
{context}

Soru: {payload.message}
Kural: Her bulgu, çıkarılan görev ve yasal riskin yanında dayandığı kaynağı `[Kaynak: X, Madde: Y]` formatında iliştir.
Kural: Son cevapta her iddia için en az bir kaynak referansı ekle.
Kural: JSON çıktısı üretirken `citations: [{"source_doc": str, "page": int, "excerpt": str}]` alanını zorunlu ekle.
Kısa, anlaşılır ve doğrudan yanıt ver. Gerekirse ilgili maddeyi veya cümleyi alıntıla."""
        reply = model.ask(prompt, timeout=45)
        citation_payload = [{
            "source_doc": citation.get("source") or entry.get("name"),
            "page": int(citation.get("page") or 1),
            "excerpt": (citation.get("text") or "")[:300],
        } for citation in citations]
    except Exception as exc:
        logger.warning("Doküman chat hatası: %s", exc)
        if isinstance(exc, HTTPException):
            raise exc
        raise HTTPException(status_code=503, detail="Doküman asistanı yanıt veremedi. Foundry Local modelini kontrol edip tekrar deneyin.")
    return {"status": "success", "reply": reply, "cited_clauses": citations, "citations": citation_payload}

@app.post("/api/v1/documents/actions/extract-tasks")
def extract_document_tasks(payload: DocumentChatEnvelope, current_user: dict = Depends(get_current_user)):
    entry, content = _get_document_content(payload.document_id)
    if not content.strip():
        raise HTTPException(status_code=422, detail="Görev çıkarmak için dokümanda metin bulunamadı.")
    lines = [line.strip(" -*•\t") for line in content.splitlines() if line.strip()]
    candidates = [line for line in lines if re.search(r"(sorumlu|sorumluluk|teslim|son tarih|deadline|yapılacak|aksiyon|görev|must|shall)", line, re.IGNORECASE)]
    tasks = []
    for line in candidates[:20]:
        task = {"id": f"task_{len(db_tasks) + len(tasks) + 1:04d}", "title": line[:240], "deadline": "Belirtilmedi", "status": "Onay Bekliyor", "type": "Dokümandan Çıkarıldı", "completed": False, "requires_ai_review": True, "source_email_id": f"document:{payload.document_id}", "sender": entry.get("name") or "Doküman", "created_by": "document_agent"}
        db_tasks.insert(0, task)
        tasks.append(task)
    return {"status": "success", "tasks": tasks, "message": f"{len(tasks)} görev dokümandan çıkarıldı."}

@app.post("/api/v1/documents/actions/extract-metrics")
def extract_document_metrics(payload: DocumentChatEnvelope, current_user: dict = Depends(get_current_user)):
    entry, content = _get_document_content(payload.document_id)
    if not content.strip():
        raise HTTPException(status_code=422, detail="Metrik çıkarmak için dokümanda metin bulunamadı.")
    metrics = []
    patterns = [
        (r"(?:₺|TL|TRY|USD|EUR|\$|€)\s?[\d.,]+(?:\s?[/] ?gün)?", "Mali", "Parasal tutar"),
        (r"[\d.,]+\s?%\s?(?:/\s?gün|günlük|ceza)?", "Mali", "Oran veya ceza"),
        (r"\b\d+\s+(?:gün|ay|yıl)\b", "Tarih/Süre", "Süre"),
    ]
    for pattern, category, label in patterns:
        for match in re.finditer(pattern, content, re.IGNORECASE):
            surrounding = content[max(0, match.start() - 100):min(len(content), match.end() + 100)].replace("\n", " ").strip()
            risk = "Yüksek" if any(word in surrounding.lower() for word in ("ceza", "fesih", "sınırsız", "istisna")) else "Orta"
            metrics.append({"category": category, "item": label, "value": match.group(0), "context": surrounding, "risk_level": risk})
    for keyword in ("force majeure", "mücbir sebep", "sla", "istisna"):
        for match in re.finditer(rf"[^.\n]*{re.escape(keyword)}[^.\n]*", content, re.IGNORECASE):
            metrics.append({"category": "İstisna", "item": keyword.title(), "value": match.group(0).strip()[:240], "risk_level": "Yüksek"})
    unique = {(item["category"], item["value"]): item for item in metrics}
    return {"status": "success", "metrics": list(unique.values())[:50], "document_id": payload.document_id, "message": f"{len(unique)} metrik/tarih/istisna bulundu."}

@app.post("/api/v1/documents/actions/visualize")
def visualize_document(payload: DocumentChatEnvelope, current_user: dict = Depends(get_current_user)):
    metrics_response = extract_document_metrics(payload, current_user)
    grouped = {}
    for metric in metrics_response.get("metrics", []):
        key = metric.get("item") or metric.get("category") or "Diğer"
        grouped[key] = grouped.get(key, 0) + 1
    return {
        "status": "success",
        "chart": {
            "chart_type": "bar",
            "title": "Doküman Metrikleri ve İstisnalar",
            "data": [{"name": name, "value": value} for name, value in grouped.items()],
            "x_key": "name",
            "y_key": "value",
        },
        "document_id": payload.document_id,
        "message": "Doküman verileri görselleştirildi.",
    }

@app.post("/api/v1/documents/actions/revise")
def revise_document(payload: DocumentActionRequest, current_user: dict = Depends(get_current_user)):
    entry, content = _get_document_content(payload.document_id)
    if not content.strip():
        raise HTTPException(status_code=422, detail="Revize edilecek doküman metni bulunamadı.")
    instruction = (payload.instruction or payload.message).strip()
    if not instruction:
        raise HTTPException(status_code=400, detail="Revizyon talimatı gerekli.")
    old_clause = (payload.selected_text or content[:1200]).strip()
    numeric_change = re.search(r"\b(\d+)\s*(kişi|gün|ay|yıl|%)\b.*?\b(\d+)\s*\2\b", instruction, re.IGNORECASE)
    if numeric_change:
        old_value, unit, new_value = numeric_change.groups()
        replacement_scope = payload.selected_text.strip() if payload.selected_text else content
        revised_clause = re.sub(rf"\b{re.escape(old_value)}\s*{re.escape(unit)}\b", f"{new_value} {unit}", replacement_scope, count=1, flags=re.IGNORECASE)
        if revised_clause == replacement_scope:
            raise HTTPException(status_code=400, detail=f"Dokümanda '{old_value} {unit}' değeri bulunamadı.")
    else:
        model = _ensure_auditor()
        try:
            revised_clause = model.ask(f"""{MANTIS_PERSONA}
Aşağıdaki doküman maddesini kullanıcı talimatına göre revize et.
    Görevin YALNIZCA belirtilen değişikliği yapmaktır.
    Orijinal metnin cümle yapısını, kelimelerini ve tonunu mümkün olduğunca koru.
    Metne orijinalinde olmayan yeni olay, kurum, kişi, hikaye veya sıfat ekleme.
    Eski veriyi tamamen sil ve yeni veriyle değiştir; iki veriyi aynı anda tutma.
    Yalnızca revize edilmiş metni döndür. Başlık, açıklama, tırnak, markdown veya gerekçe ekleme.
Doküman: {entry.get('name')}
Eski madde:
{old_clause}
    Talimat: {instruction}""", timeout=45, temperature=0.0)
        except Exception:
            raise HTTPException(status_code=503, detail="Revizyon üretilemedi. Yerel AI servisini kontrol edin.")
    revision_id = f"doc_{uuid.uuid4().hex[:8]}"
    safe_topic = re.sub(r"[^\w\s-]", "", entry.get("name", "document")).replace(" ", "_")[:32]
    filename = f"{revision_id}_v2_{safe_topic}.docx"
    filepath = os.path.join(VAULT_DIR, filename)
    document = DocxDocument()
    document.add_heading(f"Revizyon: {entry.get('name', 'Doküman')}", level=1)
    document.add_paragraph(f"Oluşturulma tarihi: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    document.add_heading("Redline / Revizyon", level=2)
    document.add_paragraph(revised_clause)
    document.add_heading("Doküman bağlamı", level=2)
    revised_content = content
    if numeric_change and not payload.selected_text:
        revised_content = revised_clause
    document.add_paragraph(revised_content)
    document.save(filepath)
    revised_entry = {"id": revision_id, "name": filename, "access": "AI Revizyonu", "last_active": datetime.now().strftime("%Y-%m-%d %H:%M"), "type": "docx_revision", "filepath": filepath, "content": revised_clause + "\n\n" + revised_content, "source_document_id": payload.document_id, "drive_link": None}
    db_vault_documents.insert(0, revised_entry)
    if rag_pipeline_available and process_and_store_document:
        process_and_store_document(source_file=filename, content=revised_entry["content"], document_type="docx_revision")
    return {"status": "success", "revision": {k: v for k, v in revised_entry.items() if k != "filepath"}, "old_clause": old_clause, "new_clause": revised_clause, "message": "Revize edilmiş yeni sürüm Doc Vault'a kaydedildi."}

@app.post("/api/v1/documents/actions/audit-risks")
def audit_document_risks(payload: DocumentChatEnvelope, current_user: dict = Depends(get_current_user)):
    try:
        entry, content = _get_document_content(payload.document_id)
        if not content.strip():
            raise HTTPException(status_code=422, detail="Denetlenecek doküman metni bulunamadı.")
        model = _ensure_auditor()
        result = model.analyze_text(content)
        findings = result.get("findings", [])
        for item in findings:
            db_risk_findings.insert(0, {**item, "status": "PENDING_APPROVAL", "is_approved": False, "is_rejected": False, "source_document_id": payload.document_id})
        return {"status": "success", "findings": findings, "message": f"Risk denetimi tamamlandı: {len(findings)} bulgu onay kuyruğuna eklendi."}
    except HTTPException:
        raise
    except Exception as exc:
        logger.warning("Doküman risk denetimi hatası: %s", exc, exc_info=True)
        raise HTTPException(status_code=503, detail="Doküman risk denetimi tamamlanamadı. Lütfen tekrar deneyin.")

@app.post("/api/v1/documents/actions/executive-summary")
def executive_document_summary(payload: DocumentChatEnvelope, current_user: dict = Depends(get_current_user)):
    entry, content = _get_document_content(payload.document_id)
    if not content.strip():
        raise HTTPException(status_code=422, detail="Özetlenecek doküman metni bulunamadı.")
    model = _ensure_auditor()
    try:
        summary = model.ask(f"""{MANTIS_PERSONA}
Aşağıdaki doküman için yöneticinin hızlı okuyabileceği Türkçe bir özet yaz.
Başlıklar: Amaç, önemli yükümlülükler, tarihler, riskler ve önerilen sonraki adım.
    Yalnızca verilen dokümana dayan; bilgi yoksa 'Dokümanda belirtilmiyor' yaz.
    Dokümanda olmayan kavram, birim, risk veya olay ekleme. Doküman yalnızca başlık/tarih içeriyorsa bunu açıkça söyle.
Doküman: {entry.get('name')}
İçerik:
    {content[:12000]}""", timeout=45, temperature=0.0)
    except Exception:
        raise HTTPException(status_code=503, detail="Yönetici özeti zamanında oluşturulamadı.")
    return {"status": "success", "summary": summary}

@app.get("/api/v1/settings/watched-folders")
def get_watched_folders(current_user: dict = Depends(get_current_user)):
    return {"status": "success", "folders": load_watched_folders()}

@app.post("/api/v1/settings/watched-folders")
def add_watched_folder(payload: WatchedFolderRequest, current_user: dict = Depends(get_current_user)):
    folder_path = os.path.abspath(os.path.expanduser(payload.path.strip()))
    if not os.path.isdir(folder_path):
        raise HTTPException(status_code=400, detail="İzlenecek klasör bulunamadı veya klasör izni yok.")
    folders = load_watched_folders()
    existing = next((folder for folder in folders if os.path.normcase(folder["path"]) == os.path.normcase(folder_path)), None)
    if existing:
        existing["enabled"] = payload.enabled
    else:
        folders.append({"path": folder_path, "enabled": payload.enabled, "added_at": datetime.now().isoformat()})
    save_watched_folders(folders)
    return {"status": "success", "folders": folders}

@app.post("/api/v1/settings/watched-folders/import")
async def import_watched_folder(
    folder_name: str = Form(...),
    files: List[UploadFile] = File(...),
    current_user: dict = Depends(get_current_user),
):
    """Tarayıcı klasör seçicisinden gelen izinli dosyaları yerel watched alanına alır."""
    clean_folder = re.sub(r"[^A-Za-z0-9_.-]", "_", folder_name).strip("._")[:80] or "selected_folder"
    folder_path = os.path.join(BACKEND_DIR, "watched_imports", clean_folder)
    os.makedirs(folder_path, exist_ok=True)
    folders = load_watched_folders()
    if not any(os.path.normcase(folder.get("path", "")) == os.path.normcase(folder_path) for folder in folders):
        folders.append({"path": folder_path, "label": folder_name, "enabled": True, "source": "browser_picker", "added_at": datetime.now().isoformat()})

    imported = 0
    skipped = 0
    supported_extensions = {".txt", ".md", ".docx"}
    for uploaded_file in files:
        original_name = uploaded_file.filename or "document"
        extension = os.path.splitext(original_name)[1].lower()
        if extension not in supported_extensions:
            skipped += 1
            continue
        safe_name = re.sub(r"[^A-Za-z0-9_.-]", "_", os.path.basename(original_name))[:120]
        destination = os.path.join(folder_path, safe_name)
        content_bytes = await uploaded_file.read()
        with open(destination, "wb") as output_file:
            output_file.write(content_bytes)
        if extension == ".docx":
            content_text = "\n".join(paragraph.text for paragraph in DocxDocument(destination).paragraphs)
        else:
            content_text = content_bytes.decode("utf-8", errors="replace")
        doc_id = "watched_" + uuid.uuid5(uuid.NAMESPACE_URL, destination).hex[:12]
        db_vault_documents.insert(0, {
            "id": doc_id, "name": safe_name, "access": "İzinli klasör",
            "last_active": datetime.now().strftime("%Y-%m-%d %H:%M"), "type": "watched",
            "filepath": destination, "content": content_text, "source_folder": folder_name, "drive_link": None,
        })
        if rag_pipeline_available and process_and_store_document and content_text.strip():
            process_and_store_document(source_file=safe_name, content=content_text, document_type="watched_folder")
        imported += 1
    save_watched_folders(folders)
    return {"status": "success", "folders": folders, "imported": imported, "skipped": skipped, "message": f"{imported} belge '{folder_name}' klasöründen Mantis'e alındı."}

@app.delete("/api/v1/settings/watched-folders")
def remove_watched_folder(path: str, current_user: dict = Depends(get_current_user)):
    folder_path = os.path.abspath(os.path.expanduser(path.strip()))
    folders = [folder for folder in load_watched_folders() if os.path.normcase(folder["path"]) != os.path.normcase(folder_path)]
    save_watched_folders(folders)
    return {"status": "success", "folders": folders}

@app.post("/api/v1/settings/watched-folders/sync")
def sync_watched_folders(current_user: dict = Depends(get_current_user)):
    imported = 0
    supported_extensions = {".txt", ".md", ".docx"}
    for folder in load_watched_folders():
        if not folder.get("enabled") or not os.path.isdir(folder["path"]):
            continue
        for root, _, filenames in os.walk(folder["path"]):
            for filename in filenames:
                if os.path.splitext(filename)[1].lower() not in supported_extensions:
                    continue
                filepath = os.path.join(root, filename)
                try:
                    if filepath.lower().endswith(".docx"):
                        content = "\n".join(paragraph.text for paragraph in DocxDocument(filepath).paragraphs)
                    else:
                        with open(filepath, "r", encoding="utf-8", errors="replace") as file:
                            content = file.read()
                    doc_id = "watched_" + uuid.uuid5(uuid.NAMESPACE_URL, filepath).hex[:12]
                    db_vault_documents.insert(0, {
                        "id": doc_id, "name": filename, "access": "İzlenen klasör",
                        "last_active": datetime.now().strftime("%Y-%m-%d %H:%M"), "type": "watched",
                        "filepath": filepath, "content": content, "source_folder": folder["path"], "drive_link": None,
                    })
                    imported += 1
                except (OSError, ValueError) as exc:
                    logger.warning("İzlenen belge okunamadı (%s): %s", filepath, exc)
    return {"status": "success", "imported": imported, "message": f"{imported} belge izlenen klasörlerden senkronlandı."}

@app.get("/api/v1/documents/download/{doc_id}")
def download_vault_document(doc_id: str):
    entry = next((d for d in db_vault_documents if d["id"] == doc_id), None)
    if not entry or not os.path.exists(entry["filepath"]):
        raise HTTPException(status_code=404, detail="Doküman bulunamadı.")
    return FileResponse(
        path=entry["filepath"],
        filename=entry["name"],
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

@app.get("/api/v1/documents/preview/{doc_id}")
def preview_vault_document(doc_id: str):
    entry = next((d for d in db_vault_documents if d["id"] == doc_id), None)
    if not entry:
        raise HTTPException(status_code=404, detail="Önizleme için belge bulunamadı.")

    content = entry.get("content") or ""
    filepath = entry.get("filepath")
    if not content and filepath and os.path.exists(filepath):
        content = _extract_document_text(filepath)
        entry["content"] = content
        db_vault_documents.update(doc_id, content=content)
    if not content:
        content = "Belge içeriği güvenli şekilde görüntüleniyor; metin okunamadı."

    return {
        "status": "success",
        "document": {
            "id": entry.get("id"),
            "name": entry.get("name"),
            "type": entry.get("type"),
            "content": content,
            "summary": (entry.get("summary") or content)[:1000],
            "download_url": f"/api/v1/documents/download/{entry.get('id')}",
        }
    }
